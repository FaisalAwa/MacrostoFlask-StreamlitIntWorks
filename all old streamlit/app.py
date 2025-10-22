import streamlit as st
from docx import Document
from io import BytesIO
import re
import time
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def ensure_spacing_before_question_tags(doc):
    """
    Ensures there's proper spacing before QUESTION NO: tags to catch all questions.
    This prevents questions from being missed when they appear directly after text.
    """
    for para in doc.paragraphs:
        # Check if paragraph contains QUESTION NO: but doesn't start with it
        if "QUESTION NO:" in para.text and not para.text.strip().startswith("QUESTION NO:"):
            # Split the text at QUESTION NO:
            parts = para.text.split("QUESTION NO:")
            if len(parts) > 1:
                # Clear the paragraph
                para.clear()
                
                # Add the text before QUESTION NO: (if any)
                if parts[0].strip():
                    para.add_run(parts[0].rstrip())
                    para.add_run().add_break()
                    para.add_run().add_break()  # Add extra line for spacing
                
                # Add QUESTION NO: and the rest
                para.add_run("QUESTION NO:" + "QUESTION NO:".join(parts[1:]))
    return doc


def apply_pagesize_macro(doc):
    for para in doc.paragraphs:
        if "QUESTION NO:" in para.text:
            para.text = re.sub(r"QUESTION NO:", "Question:", para.text)
    return doc

def remove_explanation_tags(doc):
    for para in doc.paragraphs:
        if "Explanation:" in para.text:
            para.text = para.text.replace("Explanation:", "")
    return doc

def replace_references_tag(doc): 
    for para in doc.paragraphs:
        if "References" in para.text:
            para.text = para.text.replace("References", "Reference")
    return doc
    

def shift_question_types_to_next_line(doc):
    for para in doc.paragraphs:
        if any(keyword in para.text for keyword in ["SIMULATION", "DRAG DROP", "HOTSPOT"]):
            match = re.search(r"(Question\s*:?)\s*(\d+)\s+(SIMULATION|DRAG DROP|HOTSPOT)", para.text)
            if match:
                question_part = f"{match.group(1).strip()} {match.group(2).strip()}"
                para.clear()
                para.alignment = None
                para.paragraph_format.left_indent = None
                para.paragraph_format.right_indent = None
                para.paragraph_format.first_line_indent = None
                run1 = para.add_run(question_part)
                para.add_run().add_break()
                run2 = para.add_run(match.group(3))
                run1.font.size = None
                run2.font.size = None
    return doc

def remove_map_tags(doc): 
    for para in doc.paragraphs:
        if "<map>" in para.text and "</map>" in para.text:
            cleaned = re.sub(r"<map>.*?</map>", "", para.text, flags=re.DOTALL)
            para.clear()
            para.add_run(cleaned.strip())
    return doc

def normalize_option_spacing(doc):
    # Ensure option lines like "A.     text" become "A. text" with one space
    for para in doc.paragraphs:
        txt = para.text
        if not txt:
            continue
        # Replace non-breaking spaces and tabs with regular spaces
        cleaned = txt.replace('\u00A0', ' ').replace('\t', ' ')
        match = re.match(r"^\s*([A-J])\s*\.\s*(.+)$", cleaned)
        if match:
            letter = match.group(1)
            value = match.group(2).strip()
            para.text = f"{letter}. {value}"
    return doc

def add_question_numbers(doc):
    counter = 1
    for para in doc.paragraphs:
        if re.match(r"^\s*Question\s*:?\s*$", para.text.strip()):
            para.text = f"Question: {counter}"
            counter += 1
        elif re.match(r"^\s*Question\s*:?\s*\d*\s*$", para.text.strip()):
            para.text = f"Question: {counter}"
            counter += 1
    return doc

def add_explanation_tags_if_text_present(doc):
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        if para.text.strip().startswith("Answer:"):
            explanation_lines = []
            j = i + 1
            while j < len(doc.paragraphs):
                current_para = doc.paragraphs[j]
                next_text = current_para.text.strip()
                if (next_text.startswith("Reference:") or 
                    next_text.startswith("Question:") or 
                    next_text.startswith("Topic") or
                    re.match(r"^\s*Question\s*:?\s*\d*\s*$", next_text)):
                    break
                cleaned_text = re.sub(r"<map>.*?</map>", "", next_text, flags=re.DOTALL).strip()
                if cleaned_text:
                    explanation_lines.append(cleaned_text)
                j += 1
            if explanation_lines:
                for k in range(i + 1, j):
                    doc.paragraphs[k].clear()
                if i + 1 < len(doc.paragraphs):
                    doc.paragraphs[i + 1].text = "Explanation:"
                current_para_index = i + 2
                for line in explanation_lines:
                    if current_para_index < len(doc.paragraphs):
                        doc.paragraphs[current_para_index].text = line
                        current_para_index += 1
                    else:
                        p = doc.paragraphs[current_para_index - 1]._element
                        new_para = OxmlElement("w:p")
                        run = OxmlElement("w:r")
                        text = OxmlElement("w:t")
                        text.text = line
                        run.append(text)
                        new_para.append(run)
                        p.addnext(new_para)
                        current_para_index += 1
                i = j - 1
        i += 1
    return doc


def add_line_spacing_after_question_answer(doc):
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        para_text = para.text.strip()
        
        # Handle Answer: tags - always add spacing (covers all variations)
        if para_text.startswith("Answer:"):
            # Insert an empty paragraph after this one
            p = para._element
            from docx.oxml import OxmlElement
            empty_para = OxmlElement("w:p")
            p.addnext(empty_para)
            i += 2  # Skip the newly added paragraph
        # Handle Question: X pattern (with or without HOTSPOT/SIMULATION/DRAG DROP)
        elif re.match(r"Question\s*:\s*\d+", para_text):
            # If it's a Question with HOTSPOT/SIMULATION/DRAG DROP, we need special handling
            if re.search(r"Question\s*:\s*\d+\s+(HOTSPOT|SIMULATION|DRAG DROP)", para_text):
                # Split the content: "Question: 1 HOTSPOT" becomes "Question: 1" + empty line + "HOTSPOT"
                match = re.search(r"(Question\s*:\s*\d+)\s+(HOTSPOT|SIMULATION|DRAG DROP)(.*)", para_text)
                if match:
                    question_part = match.group(1)
                    question_type = match.group(2) 
                    remaining_text = match.group(3).strip()
                    
                    # Clear and rebuild the paragraph
                    para.clear()
                    para.add_run(question_part) 
                    
                    # Add empty paragraph
                    p = para._element
                    from docx.oxml import OxmlElement
                    empty_para = OxmlElement("w:p")
                    p.addnext(empty_para)
                    
                    # Add question type paragraph  
                    type_para = OxmlElement("w:p")
                    run = OxmlElement("w:r")
                    text = OxmlElement("w:t")
                    text.text = question_type + (" " + remaining_text if remaining_text else "")
                    run.append(text)
                    type_para.append(run)
                    empty_para.addnext(type_para)
                    
                    i += 3  # Skip the paragraphs we just added
            else:
                # Regular Question: X - just add spacing
                p = para._element
                from docx.oxml import OxmlElement
                empty_para = OxmlElement("w:p")
                p.addnext(empty_para)
                i += 2
        else:
            i += 1
    
    return doc

    
def process_docx_file(uploaded_file, add_numbers=False):
    # Track macro application status and timing
    macro_status = {}
    
    doc = Document(uploaded_file)
    
    start_time = time.time()
    doc = apply_pagesize_macro(doc)
    macro_status["QUESTION NO: to Question:"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
    
    start_time = time.time()
    doc = remove_explanation_tags(doc)
    macro_status["Remove Explanation Tags"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
    
    start_time = time.time()
    doc = replace_references_tag(doc)
    macro_status["References to Reference"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
    
    start_time = time.time()
    doc = shift_question_types_to_next_line(doc)
    macro_status["Question Types to Next Line"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
    
    start_time = time.time()
    doc = remove_map_tags(doc)
    macro_status["Remove Map Tags"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}

    start_time = time.time()
    doc = normalize_option_spacing(doc)
    macro_status["Normalize Option Spacing"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
    
    if add_numbers:
        start_time = time.time()
        doc = add_question_numbers(doc)
        macro_status["Add Question Numbers"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
    else:
        macro_status["Add Question Numbers"] = {"status": "Skipped", "time": "0.00s"}
    
    start_time = time.time()
    doc = add_explanation_tags_if_text_present(doc)
    macro_status["Add Explanation Tags"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
    
    start_time = time.time()
    doc = add_line_spacing_after_question_answer(doc)
    macro_status["Add Line Spacing"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output, macro_status

# Streamlit UI
st.title("Macro Operations Via Python 🐍")

uploaded_file = st.file_uploader("Upload DOCX File", type=["docx"])
# st.text(uploaded_file.name)


with st.expander("📘 View Document Processing Steps"):
    st.text(
        "1. Ensure proper spacing before 'QUESTION NO:' tags to catch all questions.\n"
        "2. Replace 'QUESTION NO:' with standardized 'Question:' label.\n"
        "3. Remove all existing 'Explanation:' tags to avoid duplication.\n"
        "4. Replace every 'References' with 'Reference'.\n"
        "5. Add line spacing after 'Question:' and 'Answer:' tags automatically.\n"
        "6. Move question types like 'SIMULATION', 'DRAG DROP', 'HOTSPOT' to the next line.\n"
        "7. Remove all <map>...</map> XML tags embedded within paragraphs.\n"
        "8. Normalize option spacing (A. text, B. text, etc.).\n"
        "9. (Optional) Add sequential numbering to all 'Question:' tags (e.g., Question: 1, 2, ...).\n"
        "10. Insert 'Explanation:' tag only if there's valid explanation text after 'Answer:'.\n"
        "    └─ Only non-image, non-<map> text is considered.\n"
        "    └─ Detected explanation text is moved under 'Explanation:' and removed from original spot.\n"
        "11. Save and return the processed .docx file for download."
    )

if uploaded_file:
    add_numbers = st.checkbox("Add Question Numbers")

    # Create placeholders for real-time updates
    status_placeholder = st.empty()
    timer_placeholder = st.empty()
    macro_status_placeholder = st.empty()
    
    # Create table headers in advance
    with macro_status_placeholder.container():
        st.subheader("Macro Application Status:")
        col1, col2, col3 = st.columns([3, 1, 1])
        with col1:
            st.write("**Macro Name**")
        with col2:
            st.write("**Status**")
        with col3:
            st.write("**Time**")

    if st.button("Apply Macros"):
        # Initialize timer display
        start_total_time = time.time()
        timer_placeholder.info("Processing file... Time elapsed: 0.00 seconds")
        
        # Initialize status table
        current_status = {}
        
        # Define a function to update the UI in real-time
        def update_ui():
            # Update timer
            current_time = time.time() - start_total_time
            timer_placeholder.info(f"Processing file... Time elapsed: {current_time:.2f} seconds")
            
            # Update macro status table
            with macro_status_placeholder.container():
                st.subheader("Macro Application Status:")
                col1, col2, col3 = st.columns([3, 1, 1])
                with col1:
                    st.write("**Macro Name**")
                with col2:
                    st.write("**Status**")
                with col3:
                    st.write("**Time**")
                
                for macro_name, details in current_status.items():
                    col1, col2, col3 = st.columns([3, 1, 1])
                    with col1:
                        st.write(macro_name)
                    with col2:
                        if details["status"] == "Applied":
                            st.write("✅ Applied")
                        elif details["status"] == "In Progress":
                            st.write("⏳ In Progress")
                        else:
                            st.write("⏭️ Skipped")
                    with col3:
                        st.write(details["time"])
        
        # Custom process_docx_file function for real-time updates
        doc = Document(uploaded_file)
        
        # First: Ensure proper spacing before QUESTION NO: tags
        current_status["Ensure Spacing Before Questions"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = ensure_spacing_before_question_tags(doc)
        current_status["Ensure Spacing Before Questions"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Second: Replace QUESTION NO: with Question:
        # QUESTION NO: to Question:
        current_status["QUESTION NO: to Question:"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = apply_pagesize_macro(doc)
        current_status["QUESTION NO: to Question:"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Remove Explanation Tags
        current_status["Remove Explanation Tags"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = remove_explanation_tags(doc)
        current_status["Remove Explanation Tags"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # References to Reference
        current_status["References to Reference"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = replace_references_tag(doc)
        current_status["References to Reference"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Question Types to Next Line
        current_status["Question Types to Next Line"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = shift_question_types_to_next_line(doc)
        current_status["Question Types to Next Line"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Remove Map Tags
        current_status["Remove Map Tags"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = remove_map_tags(doc)
        current_status["Remove Map Tags"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()

        # Normalize Option Spacing (A. text)
        current_status["Normalize Option Spacing"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = normalize_option_spacing(doc)
        current_status["Normalize Option Spacing"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Add Question Numbers (optional)
        if add_numbers:
            current_status["Add Question Numbers"] = {"status": "In Progress", "time": "0.00s"}
            update_ui()
            start_time = time.time()
            doc = add_question_numbers(doc)
            current_status["Add Question Numbers"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        else:
            current_status["Add Question Numbers"] = {"status": "Skipped", "time": "0.00s"}
        update_ui()
        
        # Add Explanation Tags
        current_status["Add Explanation Tags"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = add_explanation_tags_if_text_present(doc)
        current_status["Add Explanation Tags"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Add Line Spacing
        current_status["Add Line Spacing"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = add_line_spacing_after_question_answer(doc)
        current_status["Add Line Spacing"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Save the document
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Final update
        total_time = time.time() - start_total_time
        status_placeholder.success(f"All macros applied successfully in {total_time:.2f} seconds!")
        
        # Show download button
        st.download_button(
            label="Download Processed File",
            data=output,
            file_name=f"{uploaded_file.name}_processed.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )






