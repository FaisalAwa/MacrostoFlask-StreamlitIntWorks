# import streamlit as st
# from docx import Document
# from io import BytesIO
# import re
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn

# def apply_pagesize_macro(doc):
#     """Replace 'QUESTION NO:' with 'Question:'"""
#     for para in doc.paragraphs:
#         if "QUESTION NO:" in para.text:
#             para.text = re.sub(r"QUESTION NO:", "Question:", para.text)
#     return doc

# def remove_explanation_tags(doc):
#     """Remove 'Explanation:' text from paragraphs"""
#     for para in doc.paragraphs:
#         if "Explanation:" in para.text:
#             para.text = para.text.replace("Explanation:", "")
#     return doc

# def replace_references_tag(doc):
#     """Replace 'References' with 'Reference'"""
#     for para in doc.paragraphs:
#         if "References" in para.text:
#             para.text = para.text.replace("References", "Reference")
#     return doc

# def shift_question_types_to_next_line(doc):
#     """Insert line break before 'SIMULATION', 'DRAG DROP', or 'HOTSPOT' words."""
#     for para in doc.paragraphs:
#         if "SIMULATION" in para.text or "DRAG DROP" in para.text or "HOTSPOT" in para.text:
#             match = re.search(r"(Question\s*:?)\s*(\d+)\s+(SIMULATION|DRAG DROP|HOTSPOT)", para.text)
#             if match:
#                 # Clean up the question part and number, removing extra spaces
#                 question_word = match.group(1).strip()
#                 question_number = match.group(2).strip()
#                 question_part = f"{question_word} {question_number}"
                
#                 # Clear the paragraph completely to remove any formatting issues
#                 para.clear()
                
#                 # Clear any paragraph-level formatting that might cause alignment issues
#                 para.alignment = None
#                 para.paragraph_format.left_indent = None
#                 para.paragraph_format.right_indent = None
#                 para.paragraph_format.first_line_indent = None
                
#                 # Add the text cleanly
#                 run1 = para.add_run(question_part)
#                 para.add_run().add_break()
#                 run2 = para.add_run(match.group(3))  # Will be "SIMULATION", "DRAG DROP", or "HOTSPOT"
                
#                 # Ensure no weird formatting on the runs
#                 run1.font.size = None
#                 run2.font.size = None
                
#     return doc

# def remove_map_tags(doc):
#     """Find and remove all <map>...</map> blocks from paragraph text."""
#     for para in doc.paragraphs:
#         if "<map>" in para.text and "</map>" in para.text:
#             cleaned = re.sub(r"<map>.*?</map>", "", para.text, flags=re.DOTALL)
#             para.clear()
#             para.add_run(cleaned.strip())
#     return doc

# def add_question_numbers(doc):
#     """Automatically number all occurrences of 'Question:' sequentially."""
#     counter = 1
#     for para in doc.paragraphs:
#         if re.match(r"^\s*Question\s*:?\s*$", para.text.strip()):
#             para.text = f"Question: {counter}"
#             counter += 1
#         elif re.match(r"^\s*Question\s*:?\s*\d*\s*$", para.text.strip()):
#             para.text = f"Question: {counter}"
#             counter += 1
#     return doc

# def add_explanation_tags_if_text_present(doc):
#     """
#     Insert 'Explanation:' only after 'Answer:' if there's meaningful explanation content
#     (excluding <map> blocks and images). Collect text until 'Reference:' or next 'Question:' tag.
#     """
#     i = 0
#     while i < len(doc.paragraphs):
#         para = doc.paragraphs[i]

#         if para.text.strip().startswith("Answer:"):
#             explanation_lines = []
#             j = i + 1

#             # Collect explanation paragraphs until 'Reference:' or next 'Question:'
#             while j < len(doc.paragraphs):
#                 current_para = doc.paragraphs[j]
#                 next_text = current_para.text.strip()

#                 # Stop if we hit Reference:, Question:, or Topic
#                 if (next_text.startswith("Reference:") or 
#                     next_text.startswith("Question:") or 
#                     next_text.startswith("Topic") or
#                     re.match(r"^\s*Question\s*:?\s*\d*\s*$", next_text)):
#                     break

#                 # Remove <map> tags and check if anything remains
#                 cleaned_text = re.sub(r"<map>.*?</map>", "", next_text, flags=re.DOTALL).strip()

#                 if cleaned_text:  # only if something remains after cleaning
#                     explanation_lines.append(cleaned_text)

#                 j += 1

#             if explanation_lines:
#                 # Clear all the explanation paragraphs first
#                 for k in range(i + 1, j):
#                     doc.paragraphs[k].clear()

#                 # Insert 'Explanation:' in the first paragraph after Answer:
#                 if i + 1 < len(doc.paragraphs):
#                     doc.paragraphs[i + 1].text = "Explanation:"
                
#                 # Insert each explanation line as separate paragraphs to preserve formatting
#                 current_para_index = i + 2
#                 for line in explanation_lines:
#                     if current_para_index < len(doc.paragraphs):
#                         doc.paragraphs[current_para_index].text = line
#                         current_para_index += 1
#                     else:
#                         # Create new paragraph for this line
#                         p = doc.paragraphs[current_para_index - 1]._element
#                         from docx.oxml import OxmlElement

#                         new_para = OxmlElement("w:p")
#                         run = OxmlElement("w:r")
#                         text = OxmlElement("w:t")
#                         text.text = line
#                         run.append(text)
#                         new_para.append(run)
#                         p.addnext(new_para)
#                         current_para_index += 1

#                 i = j - 1  # Skip already processed content

#         i += 1

#     return doc

# def add_line_spacing_after_question_answer(doc):
#     """Add line spacing after Question: and Answer: tags."""
#     i = 0
#     while i < len(doc.paragraphs):
#         para = doc.paragraphs[i]
#         para_text = para.text.strip()
        
#         # Check for Question: X pattern (with or without HOTSPOT/SIMULATION/DRAG DROP)
#         if (re.match(r"Question\s*:\s*\d+", para_text) or para_text.startswith("Answer:")):
#             # If it's a Question with HOTSPOT/SIMULATION/DRAG DROP, we need special handling
#             if re.search(r"Question\s*:\s*\d+\s+(HOTSPOT|SIMULATION|DRAG DROP)", para_text):
#                 # Split the content: "Question: 1 HOTSPOT" becomes "Question: 1" + empty line + "HOTSPOT"
#                 match = re.search(r"(Question\s*:\s*\d+)\s+(HOTSPOT|SIMULATION|DRAG DROP)(.*)", para_text)
#                 if match:
#                     question_part = match.group(1)
#                     question_type = match.group(2) 
#                     remaining_text = match.group(3).strip()
                    
#                     # Clear and rebuild the paragraph
#                     para.clear()
#                     para.add_run(question_part)
                    
#                     # Add empty paragraph
#                     p = para._element
#                     from docx.oxml import OxmlElement
#                     empty_para = OxmlElement("w:p")
#                     p.addnext(empty_para)
                    
#                     # Add question type paragraph  
#                     type_para = OxmlElement("w:p")
#                     run = OxmlElement("w:r")
#                     text = OxmlElement("w:t")
#                     text.text = question_type + (" " + remaining_text if remaining_text else "")
#                     run.append(text)
#                     type_para.append(run)
#                     empty_para.addnext(type_para)
                    
#                     i += 3  # Skip the paragraphs we just added
#             else:
#                 # Regular Question: X or Answer: - just add spacing
#                 p = para._element
#                 from docx.oxml import OxmlElement
#                 empty_para = OxmlElement("w:p")
#                 p.addnext(empty_para)
#                 i += 2
#         else:
#             i += 1
    
#     return doc

# def process_docx_file(uploaded_file, add_numbers=False, add_explanation=False, add_spacing=False):
#     doc = Document(uploaded_file)
#     doc = apply_pagesize_macro(doc)
#     doc = remove_explanation_tags(doc)
#     doc = replace_references_tag(doc)
#     if add_spacing:
#         doc = add_line_spacing_after_question_answer(doc)  # Run BEFORE shift function
#     doc = shift_question_types_to_next_line(doc)  # Run AFTER spacing
#     doc = remove_map_tags(doc)
#     if add_numbers:
#         doc = add_question_numbers(doc)
#     if add_explanation:
#         doc = add_explanation_tags_if_text_present(doc)

#     output = BytesIO()
#     doc.save(output)
#     output.seek(0)
#     return output

# # Streamlit UI
# st.title("DOCX Macro Converter")

# uploaded_file = st.file_uploader("Upload DOCX File", type=["docx"])


# with st.expander("📘 View Document Processing Steps"):
#     st.text(
#         "1. Replace 'QUESTION NO:' with standardized 'Question:' label.\n"
#         "2. Remove all existing 'Explanation:' tags to avoid duplication.\n"
#         "3. Replace every 'References' with 'Reference'.\n"
#         "4. (Optional) Add line spacing after 'Question:' and 'Answer:' tags if checkbox is selected.\n"
#         "5. Move question types like 'SIMULATION', 'DRAG DROP', 'HOTSPOT' to the next line.\n"
#         "6. Remove all <map>...</map> XML tags embedded within paragraphs.\n"
#         "7. (Optional) Add sequential numbering to all 'Question:' tags (e.g., Question: 1, 2, ...).\n"
#         "8. (Optional) Insert 'Explanation:' tag only if there's valid explanation text after 'Answer:'.\n"
#         "   └─ Only non-image, non-<map> text is considered.\n"
#         "   └─ Detected explanation text is moved under 'Explanation:' and removed from original spot.\n"
#         "9. Save and return the processed .docx file for download."
#     )


# if uploaded_file:
#     add_numbers = st.checkbox("Add Question Numbers")
#     add_explanation = st.checkbox("Add Explanation Tags")
#     add_spacing = st.checkbox("Add Line Spacing after Question & Answer")

#     if st.button("Apply Macros"):
#         processed_file = process_docx_file(
#             uploaded_file,
#             add_numbers=add_numbers,
#             add_explanation=add_explanation,
#             add_spacing=add_spacing
#         )
#         st.success("All macros applied successfully!")
#         st.download_button(
#             label="Download Processed File",
#             data=processed_file,
#             file_name="processed.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )


import streamlit as st
from docx import Document
from io import BytesIO
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def apply_pagesize_macro(doc):
    for para in doc.paragraphs:
        if "QUESTION NO:" in para.text:
            para.text = re.sub(r"QUESTION NO:", "Question:", para.text)
    return doc

def remove_explanation_tags(doc):
    for para in doc.paragraphs:
        if "Explanation:" in para.text:
            para.text = para.text.replace("Explanation:", "")
    return doc

def replace_references_tag(doc):
    for para in doc.paragraphs:
        if "References" in para.text:
            para.text = para.text.replace("References", "Reference")
    return doc

def shift_question_types_to_next_line(doc):
    for para in doc.paragraphs:
        if any(keyword in para.text for keyword in ["SIMULATION", "DRAG DROP", "HOTSPOT"]):
            match = re.search(r"(Question\s*:?)\s*(\d+)\s+(SIMULATION|DRAG DROP|HOTSPOT)", para.text)
            if match:
                question_part = f"{match.group(1).strip()} {match.group(2).strip()}"
                para.clear()
                para.alignment = None
                para.paragraph_format.left_indent = None
                para.paragraph_format.right_indent = None
                para.paragraph_format.first_line_indent = None
                run1 = para.add_run(question_part)
                para.add_run().add_break()
                run2 = para.add_run(match.group(3))
                run1.font.size = None
                run2.font.size = None
    return doc

def remove_map_tags(doc):
    for para in doc.paragraphs:
        if "<map>" in para.text and "</map>" in para.text:
            cleaned = re.sub(r"<map>.*?</map>", "", para.text, flags=re.DOTALL)
            para.clear()
            para.add_run(cleaned.strip())
    return doc

def add_question_numbers(doc):
    counter = 1
    for para in doc.paragraphs:
        if re.match(r"^\s*Question\s*:?\s*$", para.text.strip()):
            para.text = f"Question: {counter}"
            counter += 1
        elif re.match(r"^\s*Question\s*:?\s*\d*\s*$", para.text.strip()):
            para.text = f"Question: {counter}"
            counter += 1
    return doc

def add_explanation_tags_if_text_present(doc):
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        if para.text.strip().startswith("Answer:"):
            explanation_lines = []
            j = i + 1
            while j < len(doc.paragraphs):
                current_para = doc.paragraphs[j]
                next_text = current_para.text.strip()
                if (next_text.startswith("Reference:") or 
                    next_text.startswith("Question:") or 
                    next_text.startswith("Topic") or
                    re.match(r"^\s*Question\s*:?\s*\d*\s*$", next_text)):
                    break
                cleaned_text = re.sub(r"<map>.*?</map>", "", next_text, flags=re.DOTALL).strip()
                if cleaned_text:
                    explanation_lines.append(cleaned_text)
                j += 1
            if explanation_lines:
                for k in range(i + 1, j):
                    doc.paragraphs[k].clear()
                if i + 1 < len(doc.paragraphs):
                    doc.paragraphs[i + 1].text = "Explanation:"
                current_para_index = i + 2
                for line in explanation_lines:
                    if current_para_index < len(doc.paragraphs):
                        doc.paragraphs[current_para_index].text = line
                        current_para_index += 1
                    else:
                        p = doc.paragraphs[current_para_index - 1]._element
                        new_para = OxmlElement("w:p")
                        run = OxmlElement("w:r")
                        text = OxmlElement("w:t")
                        text.text = line
                        run.append(text)
                        new_para.append(run)
                        p.addnext(new_para)
                        current_para_index += 1
                i = j - 1
        i += 1
    return doc

# def add_line_spacing_after_question_answer(doc):
#     i = 0
#     while i < len(doc.paragraphs):
#         para = doc.paragraphs[i]
#         para_text = para.text.strip()
#         if (re.match(r"Question\s*:\s*\d+", para_text) or para_text.startswith("Answer:")):
#             if re.search(r"Question\s*:\s*\d+\s+(HOTSPOT|SIMULATION|DRAG DROP)", para_text):
#                 match = re.search(r"(Question\s*:\s*\d+)\s+(HOTSPOT|SIMULATION|DRAG DROP)(.*)", para_text)
#                 if match:
#                     question_part = match.group(1)
#                     question_type = match.group(2) 
#                     remaining_text = match.group(3).strip()
#                     para.clear()
#                     para.add_run(question_part)
#                     p = para._element
#                     empty_para = OxmlElement("w:p")
#                     p.addnext(empty_para)
#                     type_para = OxmlElement("w:p")
#                     run = OxmlElement("w:r")
#                     text = OxmlElement("w:t")
#                     text.text = question_type + (" " + remaining_text if remaining_text else "")
#                     run.append(text)
#                     type_para.append(run)
#                     empty_para.addnext(type_para)
#                     i += 3
#             else:
#                 p = para._element
#                 empty_para = OxmlElement("w:p")
#                 p.addnext(empty_para)
#                 i += 2
#         else:
#             i += 1
#     return doc


def add_line_spacing_after_question_answer(doc):
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        para_text = para.text.strip()
        
        # Handle Answer: tags - always add spacing (covers all variations)
        if para_text.startswith("Answer:"):
            # Insert an empty paragraph after this one
            p = para._element
            from docx.oxml import OxmlElement
            empty_para = OxmlElement("w:p")
            p.addnext(empty_para)
            i += 2  # Skip the newly added paragraph
        # Handle Question: X pattern (with or without HOTSPOT/SIMULATION/DRAG DROP)
        elif re.match(r"Question\s*:\s*\d+", para_text):
            # If it's a Question with HOTSPOT/SIMULATION/DRAG DROP, we need special handling
            if re.search(r"Question\s*:\s*\d+\s+(HOTSPOT|SIMULATION|DRAG DROP)", para_text):
                # Split the content: "Question: 1 HOTSPOT" becomes "Question: 1" + empty line + "HOTSPOT"
                match = re.search(r"(Question\s*:\s*\d+)\s+(HOTSPOT|SIMULATION|DRAG DROP)(.*)", para_text)
                if match:
                    question_part = match.group(1)
                    question_type = match.group(2) 
                    remaining_text = match.group(3).strip()
                    
                    # Clear and rebuild the paragraph
                    para.clear()
                    para.add_run(question_part)
                    
                    # Add empty paragraph
                    p = para._element
                    from docx.oxml import OxmlElement
                    empty_para = OxmlElement("w:p")
                    p.addnext(empty_para)
                    
                    # Add question type paragraph  
                    type_para = OxmlElement("w:p")
                    run = OxmlElement("w:r")
                    text = OxmlElement("w:t")
                    text.text = question_type + (" " + remaining_text if remaining_text else "")
                    run.append(text)
                    type_para.append(run)
                    empty_para.addnext(type_para)
                    
                    i += 3  # Skip the paragraphs we just added
            else:
                # Regular Question: X - just add spacing
                p = para._element
                from docx.oxml import OxmlElement
                empty_para = OxmlElement("w:p")
                p.addnext(empty_para)
                i += 2
        else:
            i += 1
    
    return doc

# def process_docx_file(uploaded_file, add_numbers=False):
#     doc = Document(uploaded_file)
#     doc = apply_pagesize_macro(doc)
#     doc = remove_explanation_tags(doc)
#     doc = replace_references_tag(doc)
#     doc = add_line_spacing_after_question_answer(doc)
#     doc = shift_question_types_to_next_line(doc)
#     doc = remove_map_tags(doc)
#     if add_numbers:
#         doc = add_question_numbers(doc)
#     doc = add_explanation_tags_if_text_present(doc)

#     output = BytesIO()
#     doc.save(output)
#     output.seek(0)
#     return output

def process_docx_file(uploaded_file, add_numbers=False):
    doc = Document(uploaded_file)
    doc = apply_pagesize_macro(doc)
    doc = remove_explanation_tags(doc)
    doc = replace_references_tag(doc)
    doc = shift_question_types_to_next_line(doc)
    doc = remove_map_tags(doc)
    if add_numbers:
        doc = add_question_numbers(doc)
    doc = add_explanation_tags_if_text_present(doc)
    # Move spacing to AFTER explanation processing
    doc = add_line_spacing_after_question_answer(doc)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Streamlit UI
st.title("Macro Operations Via Python 🐍")

uploaded_file = st.file_uploader("Upload DOCX File", type=["docx"])

with st.expander("📘 View Document Processing Steps"):
    st.text(
        "1. Replace 'QUESTION NO:' with standardized 'Question:' label.\n"
        "2. Remove all existing 'Explanation:' tags to avoid duplication.\n"
        "3. Replace every 'References' with 'Reference'.\n"
        "4. Add line spacing after 'Question:' and 'Answer:' tags automatically.\n"
        "5. Move question types like 'SIMULATION', 'DRAG DROP', 'HOTSPOT' to the next line.\n"
        "6. Remove all <map>...</map> XML tags embedded within paragraphs.\n"
        "7. (Optional) Add sequential numbering to all 'Question:' tags (e.g., Question: 1, 2, ...).\n"
        "8. Insert 'Explanation:' tag only if there's valid explanation text after 'Answer:'.\n"
        "   └─ Only non-image, non-<map> text is considered.\n"
        "   └─ Detected explanation text is moved under 'Explanation:' and removed from original spot.\n"
        "9. Save and return the processed .docx file for download."
    )

if uploaded_file:
    # add_numbers = st.checkbox("Add Question Numbers")

    if st.button("Apply Macros"):
        processed_file = process_docx_file(
            uploaded_file,
            # add_numbers=add_numbers
        )
        st.success("All macros applied successfully!")
        st.download_button(
            label="Download Processed File",
            data=processed_file,
            file_name="processed.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
