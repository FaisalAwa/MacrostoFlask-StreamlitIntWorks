import streamlit as st
from docx import Document
from io import BytesIO
import re
import time
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Pre-compile regex patterns for better performance
QUESTION_PATTERN = re.compile(r"^\s*Question\s*:?\s*\d*\s*$")
QUESTION_NUMBER_PATTERN = re.compile(r"Question\s*:\s*\d+")
QUESTION_TYPE_PATTERN = re.compile(r"(Question\s*:\s*\d+)\s+(HOTSPOT|SIMULATION|DRAG DROP)(.*)")
OPTION_PATTERN = re.compile(r"^\s*([A-J])\s*\.\s*(.+)$")
MAP_TAG_PATTERN = re.compile(r"<map>.*?</map>", flags=re.DOTALL)


def ensure_spacing_before_question_tags(doc):
    """
    Ensures there's proper spacing before QUESTION NO: tags to catch all questions.
    This prevents questions from being missed when they appear directly after text.
    """
    for para in doc.paragraphs:
        # Check if paragraph contains QUESTION NO: but doesn't start with it
        if "QUESTION NO:" in para.text and not para.text.strip().startswith("QUESTION NO:"):
            # Split the text at QUESTION NO:
            parts = para.text.split("QUESTION NO:")
            if len(parts) > 1:
                # Clear the paragraph
                para.clear()
                
                # Add the text before QUESTION NO: (if any)
                if parts[0].strip():
                    para.add_run(parts[0].rstrip())
                    para.add_run().add_break()
                    para.add_run().add_break()  # Add extra line for spacing
                
                # Add QUESTION NO: and the rest
                para.add_run("QUESTION NO:" + "QUESTION NO:".join(parts[1:]))
    return doc


def combined_text_operations(doc):
    """
    Optimized: Combine multiple simple text operations in one pass
    - Replace QUESTION NO: with Question:
    - Remove Explanation: tags
    - Replace References with Reference
    - Remove <map> tags
    """
    for para in doc.paragraphs:
        text = para.text
        if not text:
            continue
            
        modified = False
        
        # QUESTION NO: to Question:
        if "QUESTION NO:" in text:
            text = text.replace("QUESTION NO:", "Question:")
            modified = True
        
        # Remove Explanation tags
        if "Explanation:" in text:
            text = text.replace("Explanation:", "")
            modified = True
        
        # References to Reference
        if "References" in text:
            text = text.replace("References", "Reference")
            modified = True
        
        # Remove map tags
        if "<map>" in text and "</map>" in text:
            text = MAP_TAG_PATTERN.sub("", text).strip()
            modified = True
        
        if modified:
            para.text = text
    
    return doc


def shift_question_types_to_next_line(doc):
    """Move question types (SIMULATION, DRAG DROP, HOTSPOT) to next line"""
    for para in doc.paragraphs:
        if any(keyword in para.text for keyword in ["SIMULATION", "DRAG DROP", "HOTSPOT"]):
            match = re.search(r"(Question\s*:?)\s*(\d+)\s+(SIMULATION|DRAG DROP|HOTSPOT)", para.text)
            if match:
                question_part = f"{match.group(1).strip()} {match.group(2).strip()}"
                para.clear()
                para.alignment = None
                para.paragraph_format.left_indent = None
                para.paragraph_format.right_indent = None
                para.paragraph_format.first_line_indent = None
                run1 = para.add_run(question_part)
                para.add_run().add_break()
                run2 = para.add_run(match.group(3))
                run1.font.size = None
                run2.font.size = None
    return doc


def normalize_option_spacing(doc):
    """Ensure option lines like "A.     text" become "A. text" with one space"""
    for para in doc.paragraphs:
        txt = para.text
        if not txt:
            continue
        # Replace non-breaking spaces and tabs with regular spaces
        cleaned = txt.replace('\u00A0', ' ').replace('\t', ' ')
        match = OPTION_PATTERN.match(cleaned)
        if match:
            letter = match.group(1)
            value = match.group(2).strip()
            para.text = f"{letter}. {value}"
    return doc


def add_explanation_tags_if_text_present(doc):
    """
    Optimized: Add Explanation: tags where needed after Answer: tags.
    Single pass with better logic and reduced DOM manipulations.
    """
    paragraphs = list(doc.paragraphs)  # Convert to list once for better indexing
    i = 0
    
    while i < len(paragraphs):
        para = paragraphs[i]
        text = para.text.strip()
        
        if text.startswith("Answer:"):
            explanation_lines = []
            j = i + 1
            
            # Collect explanation lines
            while j < len(paragraphs):
                next_text = paragraphs[j].text.strip()
                
                # Stop conditions - check if we've reached next section
                if (next_text.startswith("Reference:") or 
                    next_text.startswith("Question:") or 
                    next_text.startswith("Topic") or
                    QUESTION_PATTERN.match(next_text)):
                    break
                
                # Clean text and collect if non-empty
                cleaned_text = MAP_TAG_PATTERN.sub("", next_text).strip()
                if cleaned_text:
                    explanation_lines.append(cleaned_text)
                
                j += 1
            
            # Only process if we found explanation text
            if explanation_lines:
                # Clear paragraphs between Answer and next section
                for k in range(i + 1, j):
                    paragraphs[k].clear()
                
                # Add "Explanation:" label
                if i + 1 < len(paragraphs):
                    paragraphs[i + 1].text = "Explanation:"
                
                # Add explanation content
                for idx, line in enumerate(explanation_lines):
                    target_idx = i + 2 + idx
                    if target_idx < len(paragraphs):
                        paragraphs[target_idx].text = line
                    else:
                        # Only create new paragraphs if absolutely needed
                        p = paragraphs[-1]._element
                        new_para = OxmlElement("w:p")
                        run = OxmlElement("w:r")
                        text_elem = OxmlElement("w:t")
                        text_elem.text = line
                        run.append(text_elem)
                        new_para.append(run)
                        p.addnext(new_para)
                        # Refresh paragraph list
                        paragraphs = list(doc.paragraphs)
                
                i = j
            else:
                i += 1
        else:
            i += 1
    
    return doc


def add_line_spacing_after_question_answer(doc):
    """
    Optimized: Add line spacing after Question: and Answer: tags.
    Process insertions in reverse to maintain indices.
    """
    paragraphs = list(doc.paragraphs)
    insertions = []  # Track where to insert empty paragraphs: (index, type, match_data)
    
    # First pass: identify where insertions are needed
    for i, para in enumerate(paragraphs):
        para_text = para.text.strip()
        
        # Check for Answer: tags
        if para_text.startswith("Answer:"):
            insertions.append((i, 'answer', None))
        
        # Check for Question: X pattern
        elif QUESTION_NUMBER_PATTERN.match(para_text):
            match = QUESTION_TYPE_PATTERN.search(para_text)
            if match:
                insertions.append((i, 'question_with_type', match))
            else:
                insertions.append((i, 'question', None))
    
    # Second pass: process insertions in reverse order to maintain correct indices
    for i, insert_type, match in reversed(insertions):
        para = paragraphs[i]
        p = para._element
        
        if insert_type == 'answer' or insert_type == 'question':
            # Simple case: just add empty paragraph after
            empty_para = OxmlElement("w:p")
            p.addnext(empty_para)
        
        elif insert_type == 'question_with_type':
            # Complex case: split question and type
            question_part = match.group(1)
            question_type = match.group(2)
            remaining_text = match.group(3).strip()
            
            # Clear and rebuild the paragraph
            para.clear()
            para.add_run(question_part)
            
            # Add empty paragraph
            empty_para = OxmlElement("w:p")
            p.addnext(empty_para)
            
            # Add question type paragraph
            type_para = OxmlElement("w:p")
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = question_type + (" " + remaining_text if remaining_text else "")
            run.append(text)
            type_para.append(run)
            empty_para.addnext(type_para)
    
    return doc


# Streamlit UI
st.title("Macro Operations Via Python 🐍")

uploaded_file = st.file_uploader("Upload DOCX File", type=["docx"])

with st.expander("📘 View Document Processing Steps"):
    st.text(
        "1. Ensure proper spacing before 'QUESTION NO:' tags to catch all questions.\n"
        "2. Combined text operations (optimized single pass):\n"
        "   - Replace 'QUESTION NO:' with standardized 'Question:' label\n"
        "   - Remove all existing 'Explanation:' tags to avoid duplication\n"
        "   - Replace every 'References' with 'Reference'\n"
        "   - Remove all <map>...</map> XML tags embedded within paragraphs\n"
        "3. Move question types like 'SIMULATION', 'DRAG DROP', 'HOTSPOT' to the next line.\n"
        "4. Normalize option spacing (A. text, B. text, etc.).\n"
        "5. Insert 'Explanation:' tag only if there's valid explanation text after 'Answer:'.\n"
        "   └─ Only non-image, non-<map> text is considered.\n"
        "   └─ Detected explanation text is moved under 'Explanation:' and removed from original spot.\n"
        "6. Add line spacing after 'Question:' and 'Answer:' tags automatically.\n"
        "7. Save and return the processed .docx file for download.\n\n"
        "⚡ Performance optimizations:\n"
        "   - Pre-compiled regex patterns for faster matching\n"
        "   - Combined multiple text operations into single pass\n"
        "   - Optimized paragraph insertion logic\n"
        "   - Reduced DOM manipulations for better speed"
    )

if uploaded_file:
    # Create placeholders for real-time updates
    status_placeholder = st.empty()
    timer_placeholder = st.empty()
    macro_status_placeholder = st.empty()
    
    # Create table headers in advance
    with macro_status_placeholder.container():
        st.subheader("Macro Application Status:")
        col1, col2, col3 = st.columns([3, 1, 1])
        with col1:
            st.write("**Macro Name**")
        with col2:
            st.write("**Status**")
        with col3:
            st.write("**Time**")

    if st.button("Apply Macros"):
        # Initialize timer display
        start_total_time = time.time()
        timer_placeholder.info("Processing file... Time elapsed: 0.00 seconds")
        
        # Initialize status table
        current_status = {}
        
        # Define a function to update the UI in real-time
        def update_ui():
            # Update timer
            current_time = time.time() - start_total_time
            timer_placeholder.info(f"Processing file... Time elapsed: {current_time:.2f} seconds")
            
            # Update macro status table
            with macro_status_placeholder.container():
                st.subheader("Macro Application Status:")
                col1, col2, col3 = st.columns([3, 1, 1])
                with col1:
                    st.write("**Macro Name**")
                with col2:
                    st.write("**Status**")
                with col3:
                    st.write("**Time**")
                
                for macro_name, details in current_status.items():
                    col1, col2, col3 = st.columns([3, 1, 1])
                    with col1:
                        st.write(macro_name)
                    with col2:
                        if details["status"] == "Applied":
                            st.write("✅ Applied")
                        elif details["status"] == "In Progress":
                            st.write("⏳ In Progress")
                        else:
                            st.write("⏭️ Skipped")
                    with col3:
                        st.write(details["time"])
        
        # Process document with real-time updates
        doc = Document(uploaded_file)
        
        # Step 1: Ensure proper spacing before QUESTION NO: tags
        current_status["Ensure Spacing Before Questions"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = ensure_spacing_before_question_tags(doc)
        current_status["Ensure Spacing Before Questions"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Step 2: Combined text operations (optimized)
        current_status["Combined Text Operations"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = combined_text_operations(doc)
        current_status["Combined Text Operations"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Step 3: Question Types to Next Line
        current_status["Question Types to Next Line"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = shift_question_types_to_next_line(doc)
        current_status["Question Types to Next Line"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Step 4: Normalize Option Spacing
        current_status["Normalize Option Spacing"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = normalize_option_spacing(doc)
        current_status["Normalize Option Spacing"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Step 5: Add Explanation Tags (optimized)
        current_status["Add Explanation Tags"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = add_explanation_tags_if_text_present(doc)
        current_status["Add Explanation Tags"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Step 6: Add Line Spacing (optimized)
        current_status["Add Line Spacing"] = {"status": "In Progress", "time": "0.00s"}
        update_ui()
        start_time = time.time()
        doc = add_line_spacing_after_question_answer(doc)
        current_status["Add Line Spacing"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
        update_ui()
        
        # Save the document
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Final update
        total_time = time.time() - start_total_time
        status_placeholder.success(f"✅ All macros applied successfully in {total_time:.2f} seconds!")
        
        # Show download button
        st.download_button(
            label="📥 Download Processed File",
            data=output,
            file_name=f"processed_{uploaded_file.name}",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )