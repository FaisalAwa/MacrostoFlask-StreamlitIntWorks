import streamlit as st
from docx import Document
from io import BytesIO
import re
import time
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import zipfile
from lxml import etree
import tempfile
import os
import shutil

# Pre-compile regex patterns for better performance
QUESTION_PATTERN = re.compile(r"^\s*Question\s*:?\s*\d*\s*$")
QUESTION_NUMBER_PATTERN = re.compile(r"Question\s*:\s*\d+")
QUESTION_TYPE_PATTERN = re.compile(r"(Question\s*:\s*\d+)\s+(HOTSPOT|SIMULATION|DRAG DROP)(.*)")
OPTION_PATTERN = re.compile(r"^\s*([A-J])\s*\.\s*(.+)$")
MAP_TAG_PATTERN = re.compile(r"<map>.*?</map>", flags=re.DOTALL)

# Valid question types for both ODT and DOCX files
VALID_QUESTION_TYPES = [
    'DRAGDROP',
    'DRAG DROP',
    'DROPDOWN',
    'HOTSPOT',
    'FILLINTHEBLANK',
    'SIMULATION',
    'POSITIONEDDRAGDROP',
    'POSITIONEDDROPDOWN'
]


# ==================== ODT FUNCTIONS ====================

def extract_valid_question_type(text):
    """
    Text se valid question type extract karta hai.
    Agar valid type nahi mila to None return karta hai.
    """
    text_upper = text.upper()
    
    for q_type in VALID_QUESTION_TYPES:
        if q_type in text_upper:
            return q_type
    
    return None


def fix_odt_question_numbers(input_file, output_file):
    """
    ODT file mein duplicate question numbers ko fix karta hai
    aur unhe proper ascending order mein arrange karta hai.
    Valid question types ko preserve karta hai.
    Format: Question: 1, Question: 2, etc. (not Question No: 01)
    Removes bracket text like [People], [Process] etc.
    """
    temp_dir = tempfile.mkdtemp()
    
    try:
        # ODT file ko extract karein (ODT ek ZIP file hai)
        with zipfile.ZipFile(input_file, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # content.xml file ko read karein
        content_file = os.path.join(temp_dir, 'content.xml')
        
        # XML parse karein
        tree = etree.parse(content_file)
        root = tree.getroot()
        
        # Namespace define karein
        namespaces = {
            'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
            'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0'
        }
        
        # Question counter
        question_counter = 1
        
        # Saare text:p elements find karein
        paragraphs = root.xpath('//text:p', namespaces=namespaces)
        
        # Pattern for both "Question No:" and "Question:"
        pattern = re.compile(r'^Question\s*(No)?:\s*\d+', re.IGNORECASE)
        
        for para in paragraphs:
            # Paragraph text extract karein
            para_text = ''.join(para.itertext())
            
            if pattern.match(para_text.strip()):
                # FIRST: Extract question type agar hai to (before any modifications)
                question_type = extract_valid_question_type(para_text)
                
                # Saare child elements remove karein
                for child in list(para):
                    para.remove(child)
                
                # SECOND: Naya text set karein with question number and type only
                # Bracket text automatically remove ho jayega kyunki hum sirf question_counter aur question_type use kar rahe hain
                if question_type:
                    para.text = f"Question: {question_counter} {question_type}"
                else:
                    para.text = f"Question: {question_counter}"
                
                para.tail = None
                
                question_counter += 1
        
        # Modified XML save karein
        tree.write(content_file, xml_declaration=True, encoding='UTF-8', pretty_print=True)
        
        # Naye ODT file mein zip karein
        with zipfile.ZipFile(output_file, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for foldername, subfolders, filenames in os.walk(temp_dir):
                for filename in filenames:
                    file_path = os.path.join(foldername, filename)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zip_ref.write(file_path, arcname)
        
        return question_counter - 1
    
    except Exception as e:
        raise Exception(f"ODT processing failed: {str(e)}")
    
    finally:
        # Temporary directory clean up
        shutil.rmtree(temp_dir, ignore_errors=True)


# ==================== DOCX FUNCTIONS ====================

def fix_docx_question_numbers_and_brackets(doc):
    """
    DOCX file mein:
    1. Question numbers ko fix karta hai (ascending order: 1, 2, 3...)
    2. Bracket text [People], [Process] etc. ko remove karta hai
    3. Valid question types ko preserve karta hai
    """
    question_counter = 1
    
    # Pattern for both "Question No:" and "Question:"
    pattern = re.compile(r'^Question\s*(No)?:\s*\d+', re.IGNORECASE)
    
    for para in doc.paragraphs:
        para_text = para.text.strip()
        
        if pattern.match(para_text):
            # Extract question type agar hai to
            question_type = extract_valid_question_type(para_text)
            
            # Clear paragraph and set new text
            para.clear()
            
            # Set new text with proper numbering (without brackets)
            if question_type:
                para.add_run(f"Question: {question_counter} {question_type}")
            else:
                para.add_run(f"Question: {question_counter}")
            
            question_counter += 1
    
    return doc


def ensure_spacing_before_question_tags(doc):
    """
    Ensures there's proper spacing before QUESTION NO: tags to catch all questions.
    This prevents questions from being missed when they appear directly after text.
    """
    for para in doc.paragraphs:
        # Check if paragraph contains QUESTION NO: but doesn't start with it
        if "QUESTION NO:" in para.text and not para.text.strip().startswith("QUESTION NO:"):
            # Split the text at QUESTION NO:
            parts = para.text.split("QUESTION NO:")
            if len(parts) > 1:
                # Clear the paragraph
                para.clear()
                
                # Add the text before QUESTION NO: (if any)
                if parts[0].strip():
                    para.add_run(parts[0].rstrip())
                    para.add_run().add_break()
                    para.add_run().add_break()  # Add extra line for spacing
                
                # Add QUESTION NO: and the rest
                para.add_run("QUESTION NO:" + "QUESTION NO:".join(parts[1:]))
    return doc


def combined_text_operations(doc):
    """
    Optimized: Combine multiple simple text operations in one pass
    - Replace QUESTION NO: with Question:
    - Remove Explanation: tags
    - Replace References with Reference
    - Remove <map> tags
    """
    for para in doc.paragraphs:
        text = para.text
        if not text:
            continue
            
        modified = False
        
        # QUESTION NO: to Question:
        if "QUESTION NO:" in text:
            text = text.replace("QUESTION NO:", "Question:")
            modified = True
        
        # Remove Explanation tags
        if "Explanation:" in text:
            text = text.replace("Explanation:", "")
            modified = True
        
        # References to Reference
        if "References" in text:
            text = text.replace("References", "Reference")
            modified = True
        
        # Remove map tags
        if "<map>" in text and "</map>" in text:
            text = MAP_TAG_PATTERN.sub("", text).strip()
            modified = True
        
        if modified:
            para.text = text
    
    return doc


def shift_question_types_to_next_line(doc):
    """Move question types (SIMULATION, DRAG DROP, HOTSPOT) to next line"""
    for para in doc.paragraphs:
        if any(keyword in para.text for keyword in ["SIMULATION", "DRAG DROP", "HOTSPOT"]):
            match = re.search(r"(Question\s*:?)\s*(\d+)\s+(SIMULATION|DRAG DROP|HOTSPOT)", para.text)
            if match:
                question_part = f"{match.group(1).strip()} {match.group(2).strip()}"
                para.clear()
                para.alignment = None
                para.paragraph_format.left_indent = None
                para.paragraph_format.right_indent = None
                para.paragraph_format.first_line_indent = None
                run1 = para.add_run(question_part)
                para.add_run().add_break()
                run2 = para.add_run(match.group(3))
                run1.font.size = None
                run2.font.size = None
    return doc


def normalize_option_spacing(doc):
    """Ensure option lines like "A.     text" become "A. text" with one space"""
    for para in doc.paragraphs:
        txt = para.text
        if not txt:
            continue
        # Replace non-breaking spaces and tabs with regular spaces
        cleaned = txt.replace('\u00A0', ' ').replace('\t', ' ')
        match = OPTION_PATTERN.match(cleaned)
        if match:
            letter = match.group(1)
            value = match.group(2).strip()
            para.text = f"{letter}. {value}"
    return doc


def add_explanation_tags_if_text_present(doc):
    """
    Optimized: Add Explanation: tags where needed after Answer: tags.
    Single pass with better logic and reduced DOM manipulations.
    """
    paragraphs = list(doc.paragraphs)  # Convert to list once for better indexing
    i = 0
    
    while i < len(paragraphs):
        para = paragraphs[i]
        text = para.text.strip()
        
        if text.startswith("Answer:"):
            explanation_lines = []
            j = i + 1
            
            # Collect explanation lines
            while j < len(paragraphs):
                next_text = paragraphs[j].text.strip()
                
                # Stop conditions - check if we've reached next section
                if (next_text.startswith("Reference:") or 
                    next_text.startswith("Question:") or 
                    next_text.startswith("Topic") or
                    QUESTION_PATTERN.match(next_text)):
                    break
                
                # Clean text and collect if non-empty
                cleaned_text = MAP_TAG_PATTERN.sub("", next_text).strip()
                if cleaned_text:
                    explanation_lines.append(cleaned_text)
                
                j += 1
            
            # Only process if we found explanation text
            if explanation_lines:
                # Clear paragraphs between Answer and next section
                for k in range(i + 1, j):
                    paragraphs[k].clear()
                
                # Add "Explanation:" label
                if i + 1 < len(paragraphs):
                    paragraphs[i + 1].text = "Explanation:"
                
                # Add explanation content
                for idx, line in enumerate(explanation_lines):
                    target_idx = i + 2 + idx
                    if target_idx < len(paragraphs):
                        paragraphs[target_idx].text = line
                    else:
                        # Only create new paragraphs if absolutely needed
                        p = paragraphs[-1]._element
                        new_para = OxmlElement("w:p")
                        run = OxmlElement("w:r")
                        text_elem = OxmlElement("w:t")
                        text_elem.text = line
                        run.append(text_elem)
                        new_para.append(run)
                        p.addnext(new_para)
                        # Refresh paragraph list
                        paragraphs = list(doc.paragraphs)
                
                i = j
            else:
                i += 1
        else:
            i += 1
    
    return doc


def add_line_spacing_after_question_answer(doc):
    """
    Optimized: Add line spacing after Question: and Answer: tags.
    Process insertions in reverse to maintain indices.
    """
    paragraphs = list(doc.paragraphs)
    insertions = []  # Track where to insert empty paragraphs: (index, type, match_data)
    
    # First pass: identify where insertions are needed
    for i, para in enumerate(paragraphs):
        para_text = para.text.strip()
        
        # Check for Answer: tags
        if para_text.startswith("Answer:"):
            insertions.append((i, 'answer', None))
        
        # Check for Question: X pattern
        elif QUESTION_NUMBER_PATTERN.match(para_text):
            match = QUESTION_TYPE_PATTERN.search(para_text)
            if match:
                insertions.append((i, 'question_with_type', match))
            else:
                insertions.append((i, 'question', None))
    
    # Second pass: process insertions in reverse order to maintain correct indices
    for i, insert_type, match in reversed(insertions):
        para = paragraphs[i]
        p = para._element
        
        if insert_type == 'answer' or insert_type == 'question':
            # Simple case: just add empty paragraph after
            empty_para = OxmlElement("w:p")
            p.addnext(empty_para)
        
        elif insert_type == 'question_with_type':
            # Complex case: split question and type
            question_part = match.group(1)
            question_type = match.group(2)
            remaining_text = match.group(3).strip()
            
            # Clear and rebuild the paragraph
            para.clear()
            para.add_run(question_part)
            
            # Add empty paragraph
            empty_para = OxmlElement("w:p")
            p.addnext(empty_para)
            
            # Add question type paragraph
            type_para = OxmlElement("w:p")
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = question_type + (" " + remaining_text if remaining_text else "")
            run.append(text)
            type_para.append(run)
            empty_para.addnext(type_para)
    
    return doc


# ==================== STREAMLIT UI ====================

st.title("Document Macro Operations 🐍")
st.markdown("### Support for both DOCX and ODT files")

uploaded_file = st.file_uploader("Upload Document File", type=["docx", "odt"])

with st.expander("📘 View Document Processing Steps"):
    st.markdown("""
    **For ODT Files:**
    1. Fix question numbering in ascending order (1, 2, 3... not 01, 02, 03)
    2. Convert "Question No:" to "Question:" format
    3. Remove bracket text like [People], [Process] etc.
    4. Preserve valid question types (HOTSPOT, SIMULATION, DRAG DROP, etc.)
    
    **For DOCX Files:**
    1. **Fix question numbering and remove brackets** (NEW!)
       - Convert "Question No:" to "Question:" format
       - Sequential numbering (1, 2, 3... not 01, 02, 03)
       - Remove bracket text like [People], [Process]
       - Preserve valid question types
    2. Ensure proper spacing before 'QUESTION NO:' tags to catch all questions
    3. Combined text operations (optimized single pass):
       - Replace 'QUESTION NO:' with standardized 'Question:' label
       - Remove all existing 'Explanation:' tags to avoid duplication
       - Replace every 'References' with 'Reference'
       - Remove all <map>...</map> XML tags embedded within paragraphs
    4. Move question types like 'SIMULATION', 'DRAG DROP', 'HOTSPOT' to the next line
    5. Normalize option spacing (A. text, B. text, etc.)
    6. Insert 'Explanation:' tag only if there's valid explanation text after 'Answer:'
       - Only non-image, non-<map> text is considered
       - Detected explanation text is moved under 'Explanation:' and removed from original spot
    7. Add line spacing after 'Question:' and 'Answer:' tags automatically
    8. Save and return the processed file for download
    
    **⚡ Performance optimizations:**
    - Pre-compiled regex patterns for faster matching
    - Combined multiple text operations into single pass
    - Optimized paragraph insertion logic
    - Reduced DOM manipulations for better speed
    """)

if uploaded_file:
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    # Create placeholders for real-time updates
    status_placeholder = st.empty()
    timer_placeholder = st.empty()
    macro_status_placeholder = st.empty()
    
    # Create table headers in advance
    with macro_status_placeholder.container():
        st.subheader("Macro Application Status:")
        col1, col2, col3 = st.columns([3, 1, 1])
        with col1:
            st.write("**Macro Name**")
        with col2:
            st.write("**Status**")
        with col3:
            st.write("**Time**")

    if st.button("Apply Macros"):
        # Initialize timer display
        start_total_time = time.time()
        timer_placeholder.info("Processing file... Time elapsed: 0.00 seconds")
        
        # Initialize status table
        current_status = {}
        
        # Define a function to update the UI in real-time
        def update_ui():
            # Update timer
            current_time = time.time() - start_total_time
            timer_placeholder.info(f"Processing file... Time elapsed: {current_time:.2f} seconds")
            
            # Update macro status table
            with macro_status_placeholder.container():
                st.subheader("Macro Application Status:")
                col1, col2, col3 = st.columns([3, 1, 1])
                with col1:
                    st.write("**Macro Name**")
                with col2:
                    st.write("**Status**")
                with col3:
                    st.write("**Time**")
                
                for macro_name, details in current_status.items():
                    col1, col2, col3 = st.columns([3, 1, 1])
                    with col1:
                        st.write(macro_name)
                    with col2:
                        if details["status"] == "Applied":
                            st.write("✅ Applied")
                        elif details["status"] == "In Progress":
                            st.write("⏳ In Progress")
                        else:
                            st.write("⏭️ Skipped")
                    with col3:
                        st.write(details["time"])
        
        # ==================== ODT FILE PROCESSING ====================
        if file_extension == "odt":
            try:
                # Step 1: Fix ODT Question Numbers
                current_status["Fix Question Numbering (ODT)"] = {"status": "In Progress", "time": "0.00s"}
                update_ui()
                start_time = time.time()
                
                # Save uploaded file to temp location
                temp_input = tempfile.NamedTemporaryFile(delete=False, suffix='.odt')
                temp_input.write(uploaded_file.read())
                temp_input.close()
                
                # Create temp output file
                temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.odt')
                temp_output.close()
                
                # Process ODT file
                questions_fixed = fix_odt_question_numbers(temp_input.name, temp_output.name)
                
                current_status["Fix Question Numbering (ODT)"] = {
                    "status": "Applied", 
                    "time": f"{(time.time() - start_time):.2f}s"
                }
                update_ui()
                
                # Read processed file
                with open(temp_output.name, 'rb') as f:
                    output_data = f.read()
                
                # Clean up temp files
                os.unlink(temp_input.name)
                os.unlink(temp_output.name)
                
                # Final update
                total_time = time.time() - start_total_time
                status_placeholder.success(
                    f"✅ ODT file processed successfully in {total_time:.2f} seconds!\n"
                    f"📊 Fixed {questions_fixed} questions"
                )
                
                # Show download button
                st.download_button(
                    label="📥 Download Processed ODT File",
                    data=output_data,
                    file_name=f"processed_{uploaded_file.name}",
                    mime="application/vnd.oasis.opendocument.text"
                )
                
            except Exception as e:
                status_placeholder.error(f"❌ Error processing ODT file: {str(e)}")
        
        # ==================== DOCX FILE PROCESSING ====================
        elif file_extension == "docx":
            try:
                doc = Document(uploaded_file)
                
                # Step 0: Fix Question Numbers and Remove Brackets (NEW FIRST STEP!)
                current_status["Fix Question Numbers & Remove Brackets"] = {"status": "In Progress", "time": "0.00s"}
                update_ui()
                start_time = time.time()
                doc = fix_docx_question_numbers_and_brackets(doc)
                current_status["Fix Question Numbers & Remove Brackets"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
                update_ui()
                
                # Step 1: Ensure proper spacing before QUESTION NO: tags
                current_status["Ensure Spacing Before Questions"] = {"status": "In Progress", "time": "0.00s"}
                update_ui()
                start_time = time.time()
                doc = ensure_spacing_before_question_tags(doc)
                current_status["Ensure Spacing Before Questions"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
                update_ui()
                
                # Step 2: Combined text operations (optimized)
                current_status["Combined Text Operations"] = {"status": "In Progress", "time": "0.00s"}
                update_ui()
                start_time = time.time()
                doc = combined_text_operations(doc)
                current_status["Combined Text Operations"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
                update_ui()
                
                # Step 3: Question Types to Next Line
                current_status["Question Types to Next Line"] = {"status": "In Progress", "time": "0.00s"}
                update_ui()
                start_time = time.time()
                doc = shift_question_types_to_next_line(doc)
                current_status["Question Types to Next Line"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
                update_ui()
                
                # Step 4: Normalize Option Spacing
                current_status["Normalize Option Spacing"] = {"status": "In Progress", "time": "0.00s"}
                update_ui()
                start_time = time.time()
                doc = normalize_option_spacing(doc)
                current_status["Normalize Option Spacing"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
                update_ui()
                
                # Step 5: Add Explanation Tags (optimized)
                current_status["Add Explanation Tags"] = {"status": "In Progress", "time": "0.00s"}
                update_ui()
                start_time = time.time()
                doc = add_explanation_tags_if_text_present(doc)
                current_status["Add Explanation Tags"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
                update_ui()
                
                # Step 6: Add Line Spacing (optimized)
                current_status["Add Line Spacing"] = {"status": "In Progress", "time": "0.00s"}
                update_ui()
                start_time = time.time()
                doc = add_line_spacing_after_question_answer(doc)
                current_status["Add Line Spacing"] = {"status": "Applied", "time": f"{(time.time() - start_time):.2f}s"}
                update_ui()
                
                # Save the document
                output = BytesIO()
                doc.save(output)
                output.seek(0)
                
                # Final update
                total_time = time.time() - start_total_time
                status_placeholder.success(f"✅ DOCX file processed successfully in {total_time:.2f} seconds!")
                
                # Show download button
                st.download_button(
                    label="📥 Download Processed DOCX File",
                    data=output,
                    file_name=f"processed_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                status_placeholder.error(f"❌ Error processing DOCX file: {str(e)}")
        
        else:
            status_placeholder.error("❌ Unsupported file format!")