# from flask import Flask, render_template, request, send_file, jsonify
# from docx import Document
# from io import BytesIO
# import re
# import time
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# import zipfile
# from lxml import etree
# import tempfile
# import os
# import shutil
# from werkzeug.utils import secure_filename

# app = Flask(__name__)
# app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
# app.config['UPLOAD_FOLDER'] = 'uploads'

# # Create uploads folder if it doesn't exist
# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# # Pre-compile regex patterns for better performance
# QUESTION_PATTERN = re.compile(r"^\s*Question\s*:?\s*\d*\s*$")
# QUESTION_NUMBER_PATTERN = re.compile(r"Question\s*:\s*\d+")
# QUESTION_TYPE_PATTERN = re.compile(r"(Question\s*:\s*\d+)\s+(HOTSPOT|SIMULATION|DRAG DROP)(.*)")
# OPTION_PATTERN = re.compile(r"^\s*([A-J])\s*\.\s*(.+)$")
# MAP_TAG_PATTERN = re.compile(r"<map>.*?</map>", flags=re.DOTALL)

# # Valid question types for both ODT and DOCX files
# VALID_QUESTION_TYPES = [
#     'DRAGDROP',
#     'DRAG DROP',
#     'DROPDOWN',
#     'HOTSPOT',
#     'FILLINTHEBLANK',
#     'SIMULATION',
#     'POSITIONEDDRAGDROP',
#     'POSITIONEDDROPDOWN'
# ]


# # ==================== ODT FUNCTIONS ====================

# def extract_valid_question_type(text):
#     """
#     Text se valid question type extract karta hai.
#     Agar valid type nahi mila to None return karta hai.
#     """
#     text_upper = text.upper()
    
#     for q_type in VALID_QUESTION_TYPES:
#         if q_type in text_upper:
#             return q_type
    
#     return None


# def fix_odt_question_numbers(input_file, output_file):
#     """
#     ODT file mein duplicate question numbers ko fix karta hai
#     aur unhe proper ascending order mein arrange karta hai.
#     Valid question types ko preserve karta hai.
#     Format: Question: 1, Question: 2, etc. (not Question No: 01)
#     Removes bracket text like [People], [Process] etc.
#     """
#     temp_dir = tempfile.mkdtemp()
    
#     try:
#         # ODT file ko extract karein (ODT ek ZIP file hai)
#         with zipfile.ZipFile(input_file, 'r') as zip_ref:
#             zip_ref.extractall(temp_dir)
        
#         # content.xml file ko read karein
#         content_file = os.path.join(temp_dir, 'content.xml')
        
#         # XML parse karein
#         tree = etree.parse(content_file)
#         root = tree.getroot()
        
#         # Namespace define karein
#         namespaces = {
#             'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
#             'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0'
#         }
        
#         # Question counter
#         question_counter = 1
        
#         # Saare text:p elements find karein
#         paragraphs = root.xpath('//text:p', namespaces=namespaces)
        
#         # Pattern for both "Question No:" and "Question:"
#         pattern = re.compile(r'^Question\s*(No)?:\s*\d+', re.IGNORECASE)
        
#         for para in paragraphs:
#             # Paragraph text extract karein
#             para_text = ''.join(para.itertext())
            
#             if pattern.match(para_text.strip()):
#                 # FIRST: Extract question type agar hai to (before any modifications)
#                 question_type = extract_valid_question_type(para_text)
                
#                 # Saare child elements remove karein
#                 for child in list(para):
#                     para.remove(child)
                
#                 # SECOND: Naya text set karein with question number and type only
#                 # Bracket text automatically remove ho jayega kyunki hum sirf question_counter aur question_type use kar rahe hain
#                 if question_type:
#                     para.text = f"Question: {question_counter} {question_type}"
#                 else:
#                     para.text = f"Question: {question_counter}"
                
#                 para.tail = None
                
#                 question_counter += 1
        
#         # Modified XML save karein
#         tree.write(content_file, xml_declaration=True, encoding='UTF-8', pretty_print=True)
        
#         # Naye ODT file mein zip karein
#         with zipfile.ZipFile(output_file, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
#             for foldername, subfolders, filenames in os.walk(temp_dir):
#                 for filename in filenames:
#                     file_path = os.path.join(foldername, filename)
#                     arcname = os.path.relpath(file_path, temp_dir)
#                     zip_ref.write(file_path, arcname)
        
#         return question_counter - 1
    
#     except Exception as e:
#         raise Exception(f"ODT processing failed: {str(e)}")
    
#     finally:
#         # Temporary directory clean up
#         shutil.rmtree(temp_dir, ignore_errors=True)


# # ==================== DOCX FUNCTIONS ====================

# def fix_docx_question_numbers_and_brackets(doc):
#     """
#     DOCX file mein:
#     1. Question numbers ko fix karta hai (ascending order: 1, 2, 3...)
#     2. Bracket text [People], [Process] etc. ko remove karta hai
#     3. Valid question types ko preserve karta hai
#     """
#     question_counter = 1
    
#     # Pattern for both "Question No:" and "Question:"
#     pattern = re.compile(r'^Question\s*(No)?:\s*\d+', re.IGNORECASE)
    
#     for para in doc.paragraphs:
#         para_text = para.text.strip()
        
#         if pattern.match(para_text):
#             # Extract question type agar hai to
#             question_type = extract_valid_question_type(para_text)
            
#             # Clear paragraph and set new text
#             para.clear()
            
#             # Set new text with proper numbering (without brackets)
#             if question_type:
#                 para.add_run(f"Question: {question_counter} {question_type}")
#             else:
#                 para.add_run(f"Question: {question_counter}")
            
#             question_counter += 1
    
#     return doc


# def ensure_spacing_before_question_tags(doc):
#     """
#     Ensures there's proper spacing before QUESTION NO: tags to catch all questions.
#     This prevents questions from being missed when they appear directly after text.
#     """
#     for para in doc.paragraphs:
#         # Check if paragraph contains QUESTION NO: but doesn't start with it
#         if "QUESTION NO:" in para.text and not para.text.strip().startswith("QUESTION NO:"):
#             # Split the text at QUESTION NO:
#             parts = para.text.split("QUESTION NO:")
#             if len(parts) > 1:
#                 # Clear the paragraph
#                 para.clear()
                
#                 # Add the text before QUESTION NO: (if any)
#                 if parts[0].strip():
#                     para.add_run(parts[0].rstrip())
#                     para.add_run().add_break()
#                     para.add_run().add_break()  # Add extra line for spacing
                
#                 # Add QUESTION NO: and the rest
#                 para.add_run("QUESTION NO:" + "QUESTION NO:".join(parts[1:]))
#     return doc


# def combined_text_operations(doc):
#     """
#     Optimized: Combine multiple simple text operations in one pass
#     - Replace QUESTION NO: with Question:
#     - Remove Explanation: tags
#     - Replace References with Reference
#     - Remove <map> tags
#     """
#     for para in doc.paragraphs:
#         text = para.text
#         if not text:
#             continue
            
#         modified = False
        
#         # QUESTION NO: to Question:
#         if "QUESTION NO:" in text:
#             text = text.replace("QUESTION NO:", "Question:")
#             modified = True
        
#         # Remove Explanation tags
#         if "Explanation:" in text:
#             text = text.replace("Explanation:", "")
#             modified = True
        
#         # References to Reference
#         if "References" in text:
#             text = text.replace("References", "Reference")
#             modified = True
        
#         # Remove map tags
#         if "<map>" in text and "</map>" in text:
#             text = MAP_TAG_PATTERN.sub("", text).strip()
#             modified = True
        
#         if modified:
#             para.text = text
    
#     return doc


# def shift_question_types_to_next_line(doc):
#     """Move question types (SIMULATION, DRAG DROP, HOTSPOT) to next line"""
#     for para in doc.paragraphs:
#         if any(keyword in para.text for keyword in ["SIMULATION", "DRAG DROP", "HOTSPOT"]):
#             match = re.search(r"(Question\s*:?)\s*(\d+)\s+(SIMULATION|DRAG DROP|HOTSPOT)", para.text)
#             if match:
#                 question_part = f"{match.group(1).strip()} {match.group(2).strip()}"
#                 para.clear()
#                 para.alignment = None
#                 para.paragraph_format.left_indent = None
#                 para.paragraph_format.right_indent = None
#                 para.paragraph_format.first_line_indent = None
#                 run1 = para.add_run(question_part)
#                 para.add_run().add_break()
#                 run2 = para.add_run(match.group(3))
#                 run1.font.size = None
#                 run2.font.size = None
#     return doc


# def normalize_option_spacing(doc):
#     """Ensure option lines like "A.     text" become "A. text" with one space"""
#     for para in doc.paragraphs:
#         txt = para.text
#         if not txt:
#             continue
#         # Replace non-breaking spaces and tabs with regular spaces
#         cleaned = txt.replace('\u00A0', ' ').replace('\t', ' ')
#         match = OPTION_PATTERN.match(cleaned)
#         if match:
#             letter = match.group(1)
#             value = match.group(2).strip()
#             para.text = f"{letter}. {value}"
#     return doc


# def add_explanation_tags_if_text_present(doc):
#     """
#     Optimized: Add Explanation: tags where needed after Answer: tags.
#     Single pass with better logic and reduced DOM manipulations.
#     """
#     paragraphs = list(doc.paragraphs)  # Convert to list once for better indexing
#     i = 0
    
#     while i < len(paragraphs):
#         para = paragraphs[i]
#         text = para.text.strip()
        
#         if text.startswith("Answer:"):
#             explanation_lines = []
#             j = i + 1
            
#             # Collect explanation lines
#             while j < len(paragraphs):
#                 next_text = paragraphs[j].text.strip()
                
#                 # Stop conditions - check if we've reached next section
#                 if (next_text.startswith("Reference:") or 
#                     next_text.startswith("Question:") or 
#                     next_text.startswith("Topic") or
#                     QUESTION_PATTERN.match(next_text)):
#                     break
                
#                 # Clean text and collect if non-empty
#                 cleaned_text = MAP_TAG_PATTERN.sub("", next_text).strip()
#                 if cleaned_text:
#                     explanation_lines.append(cleaned_text)
                
#                 j += 1
            
#             # Only process if we found explanation text
#             if explanation_lines:
#                 # Clear paragraphs between Answer and next section
#                 for k in range(i + 1, j):
#                     paragraphs[k].clear()
                
#                 # Add "Explanation:" label
#                 if i + 1 < len(paragraphs):
#                     paragraphs[i + 1].text = "Explanation:"
                
#                 # Add explanation content
#                 for idx, line in enumerate(explanation_lines):
#                     target_idx = i + 2 + idx
#                     if target_idx < len(paragraphs):
#                         paragraphs[target_idx].text = line
#                     else:
#                         # Only create new paragraphs if absolutely needed
#                         p = paragraphs[-1]._element
#                         new_para = OxmlElement("w:p")
#                         run = OxmlElement("w:r")
#                         text_elem = OxmlElement("w:t")
#                         text_elem.text = line
#                         run.append(text_elem)
#                         new_para.append(run)
#                         p.addnext(new_para)
#                         # Refresh paragraph list
#                         paragraphs = list(doc.paragraphs)
                
#                 i = j
#             else:
#                 i += 1
#         else:
#             i += 1
    
#     return doc


# def add_line_spacing_after_question_answer(doc):
#     """
#     Optimized: Add line spacing after Question: and Answer: tags.
#     Process insertions in reverse to maintain indices.
#     """
#     paragraphs = list(doc.paragraphs)
#     insertions = []  # Track where to insert empty paragraphs: (index, type, match_data)
    
#     # First pass: identify where insertions are needed
#     for i, para in enumerate(paragraphs):
#         para_text = para.text.strip()
        
#         # Check for Answer: tags
#         if para_text.startswith("Answer:"):
#             insertions.append((i, 'answer', None))
        
#         # Check for Question: X pattern
#         elif QUESTION_NUMBER_PATTERN.match(para_text):
#             match = QUESTION_TYPE_PATTERN.search(para_text)
#             if match:
#                 insertions.append((i, 'question_with_type', match))
#             else:
#                 insertions.append((i, 'question', None))
    
#     # Second pass: process insertions in reverse order to maintain correct indices
#     for i, insert_type, match in reversed(insertions):
#         para = paragraphs[i]
#         p = para._element
        
#         if insert_type == 'answer' or insert_type == 'question':
#             # Simple case: just add empty paragraph after
#             empty_para = OxmlElement("w:p")
#             p.addnext(empty_para)
        
#         elif insert_type == 'question_with_type':
#             # Complex case: split question and type
#             question_part = match.group(1)
#             question_type = match.group(2)
#             remaining_text = match.group(3).strip()
            
#             # Clear and rebuild the paragraph
#             para.clear()
#             para.add_run(question_part)
            
#             # Add empty paragraph
#             empty_para = OxmlElement("w:p")
#             p.addnext(empty_para)
            
#             # Add question type paragraph
#             type_para = OxmlElement("w:p")
#             run = OxmlElement("w:r")
#             text = OxmlElement("w:t")
#             text.text = question_type + (" " + remaining_text if remaining_text else "")
#             run.append(text)
#             type_para.append(run)
#             empty_para.addnext(type_para)
    
#     return doc


# # ==================== FLASK ROUTES ====================

# @app.route('/')
# def index():
#     return render_template('index.html')


# @app.route('/process', methods=['POST'])
# def process_document():
#     if 'file' not in request.files:
#         return jsonify({'error': 'No file uploaded'}), 400
    
#     file = request.files['file']
    
#     if file.filename == '':
#         return jsonify({'error': 'No file selected'}), 400
    
#     filename = secure_filename(file.filename)
#     file_extension = filename.split('.')[-1].lower()
    
#     if file_extension not in ['docx', 'odt']:
#         return jsonify({'error': 'Only DOCX and ODT files are supported'}), 400
    
#     start_total_time = time.time()
#     status_updates = []
    
#     try:
#         # ==================== ODT FILE PROCESSING ====================
#         if file_extension == 'odt':
#             # Save uploaded file to temp location
#             temp_input = tempfile.NamedTemporaryFile(delete=False, suffix='.odt')
#             file.save(temp_input.name)
#             temp_input.close()
            
#             # Create temp output file
#             temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.odt')
#             temp_output.close()
            
#             # Process ODT file
#             start_time = time.time()
#             status_updates.append({
#                 'name': 'Fix Question Numbering (ODT)',
#                 'status': 'in_progress',
#                 'time': '0.00s'
#             })
            
#             questions_fixed = fix_odt_question_numbers(temp_input.name, temp_output.name)
            
#             status_updates[-1] = {
#                 'name': 'Fix Question Numbering (ODT)',
#                 'status': 'completed',
#                 'time': f"{(time.time() - start_time):.2f}s"
#             }
            
#             # Read processed file
#             with open(temp_output.name, 'rb') as f:
#                 output_data = f.read()
            
#             # Clean up temp files
#             os.unlink(temp_input.name)
#             os.unlink(temp_output.name)
            
#             total_time = time.time() - start_total_time
            
#             # Save to uploads folder
#             output_filename = f"processed_{filename}"
#             output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
#             with open(output_path, 'wb') as f:
#                 f.write(output_data)
            
#             return jsonify({
#                 'success': True,
#                 'filename': output_filename,
#                 'total_time': f"{total_time:.2f}",
#                 'questions_fixed': questions_fixed,
#                 'status_updates': status_updates
#             })
        
#         # ==================== DOCX FILE PROCESSING ====================
#         elif file_extension == 'docx':
#             doc = Document(file)
            
#             # Step 0: Fix Question Numbers and Remove Brackets
#             start_time = time.time()
#             status_updates.append({
#                 'name': 'Fix Question Numbers & Remove Brackets',
#                 'status': 'in_progress',
#                 'time': '0.00s'
#             })
#             doc = fix_docx_question_numbers_and_brackets(doc)
#             status_updates[-1]['status'] = 'completed'
#             status_updates[-1]['time'] = f"{(time.time() - start_time):.2f}s"
            
#             # Step 1: Ensure proper spacing
#             start_time = time.time()
#             status_updates.append({
#                 'name': 'Ensure Spacing Before Questions',
#                 'status': 'in_progress',
#                 'time': '0.00s'
#             })
#             doc = ensure_spacing_before_question_tags(doc)
#             status_updates[-1]['status'] = 'completed'
#             status_updates[-1]['time'] = f"{(time.time() - start_time):.2f}s"
            
#             # Step 2: Combined text operations
#             start_time = time.time()
#             status_updates.append({
#                 'name': 'Combined Text Operations',
#                 'status': 'in_progress',
#                 'time': '0.00s'
#             })
#             doc = combined_text_operations(doc)
#             status_updates[-1]['status'] = 'completed'
#             status_updates[-1]['time'] = f"{(time.time() - start_time):.2f}s"
            
#             # Step 3: Question Types to Next Line
#             start_time = time.time()
#             status_updates.append({
#                 'name': 'Question Types to Next Line',
#                 'status': 'in_progress',
#                 'time': '0.00s'
#             })
#             doc = shift_question_types_to_next_line(doc)
#             status_updates[-1]['status'] = 'completed'
#             status_updates[-1]['time'] = f"{(time.time() - start_time):.2f}s"
            
#             # Step 4: Normalize Option Spacing
#             start_time = time.time()
#             status_updates.append({
#                 'name': 'Normalize Option Spacing',
#                 'status': 'in_progress',
#                 'time': '0.00s'
#             })
#             doc = normalize_option_spacing(doc)
#             status_updates[-1]['status'] = 'completed'
#             status_updates[-1]['time'] = f"{(time.time() - start_time):.2f}s"
            
#             # Step 5: Add Explanation Tags
#             start_time = time.time()
#             status_updates.append({
#                 'name': 'Add Explanation Tags',
#                 'status': 'in_progress',
#                 'time': '0.00s'
#             })
#             doc = add_explanation_tags_if_text_present(doc)
#             status_updates[-1]['status'] = 'completed'
#             status_updates[-1]['time'] = f"{(time.time() - start_time):.2f}s"
            
#             # Step 6: Add Line Spacing
#             start_time = time.time()
#             status_updates.append({
#                 'name': 'Add Line Spacing',
#                 'status': 'in_progress',
#                 'time': '0.00s'
#             })
#             doc = add_line_spacing_after_question_answer(doc)
#             status_updates[-1]['status'] = 'completed'
#             status_updates[-1]['time'] = f"{(time.time() - start_time):.2f}s"
            
#             # Save the document
#             output = BytesIO()
#             doc.save(output)
#             output.seek(0)
            
#             total_time = time.time() - start_total_time
            
#             # Save to uploads folder
#             output_filename = f"processed_{filename}"
#             output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
#             with open(output_path, 'wb') as f:
#                 f.write(output.read())
            
#             return jsonify({
#                 'success': True,
#                 'filename': output_filename,
#                 'total_time': f"{total_time:.2f}",
#                 'status_updates': status_updates
#             })
    
#     except Exception as e:
#         return jsonify({'error': str(e)}), 500


# @app.route('/download/<filename>')
# def download_file(filename):
#     try:
#         file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#         return send_file(file_path, as_attachment=True)
#     except Exception as e:
#         return jsonify({'error': str(e)}), 404


# if __name__ == '__main__':
#     app.run(debug=True, host='0.0.0.0', port=5000)




from flask import Flask, render_template, request, send_file, jsonify, Response
from docx import Document
from io import BytesIO
import re
import time
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import zipfile
from lxml import etree
import tempfile
import os
import shutil
from werkzeug.utils import secure_filename
import json
import queue
import threading

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'

# Create uploads folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Queue for real-time status updates
status_queue = queue.Queue()

# Pre-compile regex patterns for better performance
QUESTION_PATTERN = re.compile(r"^\s*Question\s*:?\s*\d*\s*$")
QUESTION_NUMBER_PATTERN = re.compile(r"Question\s*:\s*\d+")
QUESTION_TYPE_PATTERN = re.compile(r"(Question\s*:\s*\d+)\s+(HOTSPOT|SIMULATION|DRAG DROP)(.*)")
OPTION_PATTERN = re.compile(r"^\s*([A-J])\s*\.\s*(.+)$")
MAP_TAG_PATTERN = re.compile(r"<map>.*?</map>", flags=re.DOTALL)

# Valid question types for both ODT and DOCX files
VALID_QUESTION_TYPES = [
    'DRAGDROP',
    'DRAG DROP',
    'DROPDOWN',
    'HOTSPOT',
    'FILLINTHEBLANK',
    'SIMULATION',
    'POSITIONEDDRAGDROP',
    'POSITIONEDDROPDOWN'
]


# ==================== ODT FUNCTIONS ====================

def extract_valid_question_type(text):
    """
    Text se valid question type extract karta hai.
    Agar valid type nahi mila to None return karta hai.
    """
    text_upper = text.upper()
    
    for q_type in VALID_QUESTION_TYPES:
        if q_type in text_upper:
            return q_type
    
    return None


def fix_odt_question_numbers(input_file, output_file):
    """
    ODT file mein duplicate question numbers ko fix karta hai
    aur unhe proper ascending order mein arrange karta hai.
    Valid question types ko preserve karta hai.
    Format: Question: 1, Question: 2, etc. (not Question No: 01)
    Removes bracket text like [People], [Process] etc.
    """
    temp_dir = tempfile.mkdtemp()
    
    try:
        # ODT file ko extract karein (ODT ek ZIP file hai)
        with zipfile.ZipFile(input_file, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # content.xml file ko read karein
        content_file = os.path.join(temp_dir, 'content.xml')
        
        # XML parse karein
        tree = etree.parse(content_file)
        root = tree.getroot()
        
        # Namespace define karein
        namespaces = {
            'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
            'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0'
        }
        
        # Question counter
        question_counter = 1
        
        # Saare text:p elements find karein
        paragraphs = root.xpath('//text:p', namespaces=namespaces)
        
        # Pattern for both "Question No:" and "Question:"
        pattern = re.compile(r'^Question\s*(No)?:\s*\d+', re.IGNORECASE)
        
        for para in paragraphs:
            # Paragraph text extract karein
            para_text = ''.join(para.itertext())
            
            if pattern.match(para_text.strip()):
                # FIRST: Extract question type agar hai to (before any modifications)
                question_type = extract_valid_question_type(para_text)
                
                # Saare child elements remove karein
                for child in list(para):
                    para.remove(child)
                
                # SECOND: Naya text set karein with question number and type only
                # Bracket text automatically remove ho jayega kyunki hum sirf question_counter aur question_type use kar rahe hain
                if question_type:
                    para.text = f"Question: {question_counter} {question_type}"
                else:
                    para.text = f"Question: {question_counter}"
                
                para.tail = None
                
                question_counter += 1
        
        # Modified XML save karein
        tree.write(content_file, xml_declaration=True, encoding='UTF-8', pretty_print=True)
        
        # Naye ODT file mein zip karein
        with zipfile.ZipFile(output_file, 'w', zipfile.ZIP_DEFLATED) as zip_ref:
            for foldername, subfolders, filenames in os.walk(temp_dir):
                for filename in filenames:
                    file_path = os.path.join(foldername, filename)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zip_ref.write(file_path, arcname)
        
        return question_counter - 1
    
    except Exception as e:
        raise Exception(f"ODT processing failed: {str(e)}")
    
    finally:
        # Temporary directory clean up
        shutil.rmtree(temp_dir, ignore_errors=True)


# ==================== DOCX FUNCTIONS ====================

def fix_docx_question_numbers_and_brackets(doc):
    """
    DOCX file mein:
    1. Question numbers ko fix karta hai (ascending order: 1, 2, 3...)
    2. Bracket text [People], [Process] etc. ko remove karta hai
    3. Valid question types ko preserve karta hai
    """
    question_counter = 1
    
    # Pattern for both "Question No:" and "Question:"
    pattern = re.compile(r'^Question\s*(No)?:\s*\d+', re.IGNORECASE)
    
    for para in doc.paragraphs:
        para_text = para.text.strip()
        
        if pattern.match(para_text):
            # Extract question type agar hai to
            question_type = extract_valid_question_type(para_text)
            
            # Clear paragraph and set new text
            para.clear()
            
            # Set new text with proper numbering (without brackets)
            if question_type:
                para.add_run(f"Question: {question_counter} {question_type}")
            else:
                para.add_run(f"Question: {question_counter}")
            
            question_counter += 1
    
    return doc


def ensure_spacing_before_question_tags(doc):
    """
    Ensures there's proper spacing before QUESTION NO: tags to catch all questions.
    This prevents questions from being missed when they appear directly after text.
    """
    for para in doc.paragraphs:
        # Check if paragraph contains QUESTION NO: but doesn't start with it
        if "QUESTION NO:" in para.text and not para.text.strip().startswith("QUESTION NO:"):
            # Split the text at QUESTION NO:
            parts = para.text.split("QUESTION NO:")
            if len(parts) > 1:
                # Clear the paragraph
                para.clear()
                
                # Add the text before QUESTION NO: (if any)
                if parts[0].strip():
                    para.add_run(parts[0].rstrip())
                    para.add_run().add_break()
                    para.add_run().add_break()  # Add extra line for spacing
                
                # Add QUESTION NO: and the rest
                para.add_run("QUESTION NO:" + "QUESTION NO:".join(parts[1:]))
    return doc


def combined_text_operations(doc):
    """
    Optimized: Combine multiple simple text operations in one pass
    - Replace QUESTION NO: with Question:
    - Remove Explanation: tags
    - Replace References with Reference
    - Remove <map> tags
    """
    for para in doc.paragraphs:
        text = para.text
        if not text:
            continue
            
        modified = False
        
        # QUESTION NO: to Question:
        if "QUESTION NO:" in text:
            text = text.replace("QUESTION NO:", "Question:")
            modified = True
        
        # Remove Explanation tags
        if "Explanation:" in text:
            text = text.replace("Explanation:", "")
            modified = True
        
        # References to Reference
        if "References" in text:
            text = text.replace("References", "Reference")
            modified = True
        
        # Remove map tags
        if "<map>" in text and "</map>" in text:
            text = MAP_TAG_PATTERN.sub("", text).strip()
            modified = True
        
        if modified:
            para.text = text
    
    return doc


def shift_question_types_to_next_line(doc):
    """Move question types (SIMULATION, DRAG DROP, HOTSPOT) to next line"""
    for para in doc.paragraphs:
        if any(keyword in para.text for keyword in ["SIMULATION", "DRAG DROP", "HOTSPOT"]):
            match = re.search(r"(Question\s*:?)\s*(\d+)\s+(SIMULATION|DRAG DROP|HOTSPOT)", para.text)
            if match:
                question_part = f"{match.group(1).strip()} {match.group(2).strip()}"
                para.clear()
                para.alignment = None
                para.paragraph_format.left_indent = None
                para.paragraph_format.right_indent = None
                para.paragraph_format.first_line_indent = None
                run1 = para.add_run(question_part)
                para.add_run().add_break()
                run2 = para.add_run(match.group(3))
                run1.font.size = None
                run2.font.size = None
    return doc


def normalize_option_spacing(doc):
    """Ensure option lines like "A.     text" become "A. text" with one space"""
    for para in doc.paragraphs:
        txt = para.text
        if not txt:
            continue
        # Replace non-breaking spaces and tabs with regular spaces
        cleaned = txt.replace('\u00A0', ' ').replace('\t', ' ')
        match = OPTION_PATTERN.match(cleaned)
        if match:
            letter = match.group(1)
            value = match.group(2).strip()
            para.text = f"{letter}. {value}"
    return doc


def add_explanation_tags_if_text_present(doc):
    """
    Optimized: Add Explanation: tags where needed after Answer: tags.
    Single pass with better logic and reduced DOM manipulations.
    """
    paragraphs = list(doc.paragraphs)  # Convert to list once for better indexing
    i = 0
    
    while i < len(paragraphs):
        para = paragraphs[i]
        text = para.text.strip()
        
        if text.startswith("Answer:"):
            explanation_lines = []
            j = i + 1
            
            # Collect explanation lines
            while j < len(paragraphs):
                next_text = paragraphs[j].text.strip()
                
                # Stop conditions - check if we've reached next section
                if (next_text.startswith("Reference:") or 
                    next_text.startswith("Question:") or 
                    next_text.startswith("Topic") or
                    QUESTION_PATTERN.match(next_text)):
                    break
                
                # Clean text and collect if non-empty
                cleaned_text = MAP_TAG_PATTERN.sub("", next_text).strip()
                if cleaned_text:
                    explanation_lines.append(cleaned_text)
                
                j += 1
            
            # Only process if we found explanation text
            if explanation_lines:
                # Clear paragraphs between Answer and next section
                for k in range(i + 1, j):
                    paragraphs[k].clear()
                
                # Add "Explanation:" label
                if i + 1 < len(paragraphs):
                    paragraphs[i + 1].text = "Explanation:"
                
                # Add explanation content
                for idx, line in enumerate(explanation_lines):
                    target_idx = i + 2 + idx
                    if target_idx < len(paragraphs):
                        paragraphs[target_idx].text = line
                    else:
                        # Only create new paragraphs if absolutely needed
                        p = paragraphs[-1]._element
                        new_para = OxmlElement("w:p")
                        run = OxmlElement("w:r")
                        text_elem = OxmlElement("w:t")
                        text_elem.text = line
                        run.append(text_elem)
                        new_para.append(run)
                        p.addnext(new_para)
                        # Refresh paragraph list
                        paragraphs = list(doc.paragraphs)
                
                i = j
            else:
                i += 1
        else:
            i += 1
    
    return doc


def add_line_spacing_after_question_answer(doc):
    """
    Optimized: Add line spacing after Question: and Answer: tags.
    Process insertions in reverse to maintain indices.
    """
    paragraphs = list(doc.paragraphs)
    insertions = []  # Track where to insert empty paragraphs: (index, type, match_data)
    
    # First pass: identify where insertions are needed
    for i, para in enumerate(paragraphs):
        para_text = para.text.strip()
        
        # Check for Answer: tags
        if para_text.startswith("Answer:"):
            insertions.append((i, 'answer', None))
        
        # Check for Question: X pattern
        elif QUESTION_NUMBER_PATTERN.match(para_text):
            match = QUESTION_TYPE_PATTERN.search(para_text)
            if match:
                insertions.append((i, 'question_with_type', match))
            else:
                insertions.append((i, 'question', None))
    
    # Second pass: process insertions in reverse order to maintain correct indices
    for i, insert_type, match in reversed(insertions):
        para = paragraphs[i]
        p = para._element
        
        if insert_type == 'answer' or insert_type == 'question':
            # Simple case: just add empty paragraph after
            empty_para = OxmlElement("w:p")
            p.addnext(empty_para)
        
        elif insert_type == 'question_with_type':
            # Complex case: split question and type
            question_part = match.group(1)
            question_type = match.group(2)
            remaining_text = match.group(3).strip()
            
            # Clear and rebuild the paragraph
            para.clear()
            para.add_run(question_part)
            
            # Add empty paragraph
            empty_para = OxmlElement("w:p")
            p.addnext(empty_para)
            
            # Add question type paragraph
            type_para = OxmlElement("w:p")
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = question_type + (" " + remaining_text if remaining_text else "")
            run.append(text)
            type_para.append(run)
            empty_para.addnext(type_para)
    
    return doc


# ==================== FLASK ROUTES ====================

def send_status_update(update_data):
    """Send status update to the queue for SSE"""
    status_queue.put(json.dumps(update_data))


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/status-stream')
def status_stream():
    """Server-Sent Events endpoint for real-time updates"""
    def generate():
        while True:
            try:
                # Get update from queue (timeout to check if still connected)
                update = status_queue.get(timeout=30)
                yield f"data: {update}\n\n"
            except queue.Empty:
                # Send heartbeat to keep connection alive
                yield f"data: {json.dumps({'type': 'heartbeat'})}\n\n"
    
    return Response(generate(), mimetype='text/event-stream')


@app.route('/process', methods=['POST'])
def process_document():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    filename = secure_filename(file.filename)
    file_extension = filename.split('.')[-1].lower()
    
    if file_extension not in ['docx', 'odt']:
        return jsonify({'error': 'Only DOCX and ODT files are supported'}), 400
    
    # Save file temporarily
    temp_file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{filename}")
    file.save(temp_file_path)
    
    # Start processing in background thread
    def process_in_background():
        start_total_time = time.time()
        
        try:
            # Clear the queue
            while not status_queue.empty():
                try:
                    status_queue.get_nowait()
                except queue.Empty:
                    break
            
            # Send initial status
            send_status_update({
                'type': 'start',
                'message': 'Processing started...'
            })
            
            # ==================== ODT FILE PROCESSING ====================
            if file_extension == 'odt':
                # Send status update
                send_status_update({
                    'type': 'status',
                    'name': 'Fix Question Numbering (ODT)',
                    'status': 'in_progress',
                    'time': '0.00s'
                })
                
                # Create temp output file
                temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.odt')
                temp_output.close()
                
                # Process ODT file
                start_time = time.time()
                questions_fixed = fix_odt_question_numbers(temp_file_path, temp_output.name)
                
                send_status_update({
                    'type': 'status',
                    'name': 'Fix Question Numbering (ODT)',
                    'status': 'completed',
                    'time': f"{(time.time() - start_time):.2f}s"
                })
                
                # Read processed file
                with open(temp_output.name, 'rb') as f:
                    output_data = f.read()
                
                # Clean up temp files
                os.unlink(temp_file_path)
                os.unlink(temp_output.name)
                
                total_time = time.time() - start_total_time
                
                # Save to uploads folder
                output_filename = f"processed_{filename}"
                output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                with open(output_path, 'wb') as f:
                    f.write(output_data)
                
                send_status_update({
                    'type': 'complete',
                    'filename': output_filename,
                    'total_time': f"{total_time:.2f}",
                    'questions_fixed': questions_fixed
                })
            
            # ==================== DOCX FILE PROCESSING ====================
            elif file_extension == 'docx':
                doc = Document(temp_file_path)
                
                # Step 0: Fix Question Numbers and Remove Brackets
                send_status_update({
                    'type': 'status',
                    'name': 'Fix Question Numbers & Remove Brackets',
                    'status': 'in_progress',
                    'time': '0.00s'
                })
                start_time = time.time()
                doc = fix_docx_question_numbers_and_brackets(doc)
                send_status_update({
                    'type': 'status',
                    'name': 'Fix Question Numbers & Remove Brackets',
                    'status': 'completed',
                    'time': f"{(time.time() - start_time):.2f}s"
                })
                
                # Step 1: Ensure proper spacing
                send_status_update({
                    'type': 'status',
                    'name': 'Ensure Spacing Before Questions',
                    'status': 'in_progress',
                    'time': '0.00s'
                })
                start_time = time.time()
                doc = ensure_spacing_before_question_tags(doc)
                send_status_update({
                    'type': 'status',
                    'name': 'Ensure Spacing Before Questions',
                    'status': 'completed',
                    'time': f"{(time.time() - start_time):.2f}s"
                })
                
                # Step 2: Combined text operations
                send_status_update({
                    'type': 'status',
                    'name': 'Combined Text Operations',
                    'status': 'in_progress',
                    'time': '0.00s'
                })
                start_time = time.time()
                doc = combined_text_operations(doc)
                send_status_update({
                    'type': 'status',
                    'name': 'Combined Text Operations',
                    'status': 'completed',
                    'time': f"{(time.time() - start_time):.2f}s"
                })
                
                # Step 3: Question Types to Next Line
                send_status_update({
                    'type': 'status',
                    'name': 'Question Types to Next Line',
                    'status': 'in_progress',
                    'time': '0.00s'
                })
                start_time = time.time()
                doc = shift_question_types_to_next_line(doc)
                send_status_update({
                    'type': 'status',
                    'name': 'Question Types to Next Line',
                    'status': 'completed',
                    'time': f"{(time.time() - start_time):.2f}s"
                })
                
                # Step 4: Normalize Option Spacing
                send_status_update({
                    'type': 'status',
                    'name': 'Normalize Option Spacing',
                    'status': 'in_progress',
                    'time': '0.00s'
                })
                start_time = time.time()
                doc = normalize_option_spacing(doc)
                send_status_update({
                    'type': 'status',
                    'name': 'Normalize Option Spacing',
                    'status': 'completed',
                    'time': f"{(time.time() - start_time):.2f}s"
                })
                
                # Step 5: Add Explanation Tags
                send_status_update({
                    'type': 'status',
                    'name': 'Add Explanation Tags',
                    'status': 'in_progress',
                    'time': '0.00s'
                })
                start_time = time.time()
                doc = add_explanation_tags_if_text_present(doc)
                send_status_update({
                    'type': 'status',
                    'name': 'Add Explanation Tags',
                    'status': 'completed',
                    'time': f"{(time.time() - start_time):.2f}s"
                })
                
                # Step 6: Add Line Spacing
                send_status_update({
                    'type': 'status',
                    'name': 'Add Line Spacing',
                    'status': 'in_progress',
                    'time': '0.00s'
                })
                start_time = time.time()
                doc = add_line_spacing_after_question_answer(doc)
                send_status_update({
                    'type': 'status',
                    'name': 'Add Line Spacing',
                    'status': 'completed',
                    'time': f"{(time.time() - start_time):.2f}s"
                })
                
                # Save the document
                output = BytesIO()
                doc.save(output)
                output.seek(0)
                
                total_time = time.time() - start_total_time
                
                # Save to uploads folder
                output_filename = f"processed_{filename}"
                output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                with open(output_path, 'wb') as f:
                    f.write(output.read())
                
                # Clean up temp file
                os.unlink(temp_file_path)
                
                send_status_update({
                    'type': 'complete',
                    'filename': output_filename,
                    'total_time': f"{total_time:.2f}"
                })
        
        except Exception as e:
            send_status_update({
                'type': 'error',
                'message': str(e)
            })
            # Clean up temp file on error
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
    
    # Start background thread
    thread = threading.Thread(target=process_in_background)
    thread.daemon = True
    thread.start()
    
    return jsonify({'success': True, 'message': 'Processing started'})


@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        return jsonify({'error': str(e)}), 404


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)